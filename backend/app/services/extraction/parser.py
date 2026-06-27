"""Document parsers for Excel and Word files."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def parse_excel(
    file_path: Path,
    column_mapping: dict[str, str],
    sheet_name: str | None = None,
    header_row: int = 1,
    ner_columns: list[str] | None = None,
) -> list[dict[str, Any]]:
    """Parse Excel file and map columns to ontology property IRIs.

    Args:
        file_path: Path to .xlsx file
        column_mapping: {column_header: property_iri}
        sheet_name: Target sheet (default: first)
        header_row: Row containing headers (1-indexed)
        ner_columns: 自由文本列白名单（008 US3，FR-007/008）。命中列原文不直接落 IRION
            属性键，而暂存于本行 ``__freetext__``（``{header: 原文}``，键集 = ner_columns ∩
            实际表头），供管线本地 NER 富化。``None`` 时无 ``__freetext__``、行为与改造前
            完全一致（向后兼容）。自由文本列优先于结构化映射（同列只暂存、不落属性）。

    Returns:
        List of dicts with property IRI keys（自由文本暂存于临时 ``__freetext__`` 键）。
    """
    import openpyxl

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    rows = list(ws.iter_rows(min_row=header_row, values_only=False))
    if not rows:
        return []

    header_cells = rows[0]
    headers = [str(cell.value).strip() if cell.value else "" for cell in header_cells]
    ner_set = set(ner_columns or [])

    # 列分流：自由文本列优先暂存（不落属性）；其余命中映射的列落 IRI 键。
    col_to_prop: dict[int, str] = {}
    col_to_freetext: dict[int, str] = {}
    for idx, header in enumerate(headers):
        if header in ner_set:
            col_to_freetext[idx] = header
        elif header in column_mapping:
            col_to_prop[idx] = column_mapping[header]

    results = []
    for row in rows[1:]:
        values = [cell.value for cell in row]
        if not any(values):
            continue
        entity = {}
        for col_idx, prop_iri in col_to_prop.items():
            val = values[col_idx] if col_idx < len(values) else None
            if val is not None:
                entity[prop_iri] = val
        freetext = {}
        for col_idx, header in col_to_freetext.items():
            val = values[col_idx] if col_idx < len(values) else None
            if val is not None and str(val).strip():
                freetext[header] = val
        if freetext:
            entity["__freetext__"] = freetext      # 临时暂存键，富化后由管线清除
        if entity:
            results.append(entity)

    wb.close()
    return results


def parse_word(
    file_path: Path,
    column_mapping: dict[str, str] | None = None,
) -> list[dict[str, Any]]:
    """Parse Word document, extracting structured tables and paragraphs.

    Returns list of extracted sections, each with type (table/paragraph) and content.

    Args:
        file_path: Path to .docx file
        column_mapping: {表头文本: property_iri}——表格行单元格键的确定性映射，
            与 ``parse_excel`` 对齐（FR-004，research R6）：表头命中映射→以 IRI 为键，
            未命中→原表头键保留（容忍未映射列）；``None`` 时全部沿用原表头键
            （向后兼容）。Word 表头→IRI 走纯本地确定性映射，不再依赖云端 LLM。
    """
    import docx

    doc = docx.Document(file_path)
    sections = []

    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            row_data = {}
            for idx, cell in enumerate(row.cells):
                if idx < len(headers) and headers[idx]:
                    header = headers[idx]
                    key = (column_mapping[header]
                           if column_mapping and header in column_mapping else header)
                    row_data[key] = cell.text.strip()
            if row_data:
                sections.append({"type": "table_row", "content": row_data})

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append({"type": "paragraph", "content": text, "style": para.style.name})

    return sections
