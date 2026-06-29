"""文档标注服务：读取上传文件 → 三阶段 NER 标注 → 前端可渲染结构 + 属性三元组。

阶段一：GLiNER 以本体种子标签集（数据属性 domain 类 + 根/近根类，30-40 量级）召回实体 span；
阶段二：每个 span 的**上下文窗口文本**经本地嵌入余弦匹配到**最具体**的本体类
（``ontology_typer.type_spans``，阈值 0.50，低于则丢弃）。span 的 label/className 取匹配
类标签，score 取余弦相似度，前端据此按类着色、tooltip 显示「实体类型 · 置信度」；
阶段三：对每个已归类实体，以其本体类的数据属性标签集跑 GLiNER 抽取属性值 → 三元组
（``_extract_property_triples``），同时供标注展示与候选审核。

GLiNER 仅负责定界（找 span 边界），其给出的种子标签被丢弃，最终类别由嵌入归类决定。

Word → tiptap JSON（ProseMirror doc，entity mark 内嵌 span）
Excel → 结构化行数据，每个单元格附 entity annotations

性能关键：整篇文档的所有文本段先收集成一个列表，**一次** GLiNER batch 推理 + **一次**
嵌入归类，再把 span 分发回各段。逐段调用在 CPU 上对几百段会到分钟级（请求超时）。

返回 ``(content, warnings, triples)``：warnings 携带 get_class_properties domain bug 的记录，
供抽取任务以告警上浮（用户：在本特性中仅记录、不修复，根因修复留待后续 feature）。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Callable

from app.config import settings
from app.services.extraction.ontology_typer import (
    GET_CLASS_PROPERTIES_DOMAIN_BUG,
    _is_non_entity_span,
    predict_segment_classes,
    seed_labels,
    type_spans,
)

logger = logging.getLogger(__name__)


def _get_extractor():
    from app.services.extraction.gliner_extractor import get_gliner_extractor

    return get_gliner_extractor()


def _extract_spans_batch(
    extractor, texts: list[str], labels: list[str]
) -> list[list[dict]]:
    """阶段一：批量召回整篇文档所有文本段的候选 span（一次 GLiNER 推理）。

    返回与 ``texts`` 等长的 span 列表的列表；第 i 项为 ``texts[i]`` 的
    ``[{start, end, text, label, score}]``（label 为种子类，稍后被归类覆盖）。
    提取器不可用 → 等长空列表列表（优雅降级，结构化主路径零回归）。
    """
    if not extractor or not extractor.is_available() or not texts or not labels:
        return [[] for _ in texts]
    return extractor.extract_batch_with_spans(texts, labels)


_CONTEXT_WINDOW = 15

# 双程归类：Pass-1 高置信直接匹配阈值（raw text，无上下文）。
_DIRECT_MATCH_THRESHOLD = 0.85
# 短 span 字符上限：span text ≤ 此长度时仅走 raw-text Pass-1，不拼上下文。
_SHORT_SPAN_MAX_CJK = 3
_SHORT_SPAN_MAX_CHARS = 4


def _is_short_span(text: str) -> bool:
    """判断 span 是否为短文本（上下文会淹没其语义）。"""
    cjk = sum(1 for ch in text if '一' <= ch <= '鿿')
    if cjk > 0:
        return cjk <= _SHORT_SPAN_MAX_CJK
    return len(text) <= _SHORT_SPAN_MAX_CHARS


def _span_with_context(
    seg_text: str, start: int, end: int, window: int = _CONTEXT_WINDOW
) -> str:
    """拼接 span 前后上下文窗口，为嵌入提供语义锚点。

    对无语义内容的 span（如 "HRS-1234"），周围上下文（"原料药…临床备样"）
    使嵌入向量携带正确语义方向，避免随机命中错误本体类。
    """
    prefix = seg_text[max(0, start - window):start]
    suffix = seg_text[end:end + window]
    span_text = seg_text[start:end]
    return f"{prefix}{span_text}{suffix}".strip()


def _type_and_filter_spans(
    all_spans: list[list[dict]],
    segment_texts: list[str],
    engine,
    class_iris: set[str] | None = None,
) -> list[list[dict]]:
    """阶段二：把每段的候选 span 按上下文窗口文本归类到最具体的本体类，低于阈值丢弃。

    两层缩窄（仅文档分类成功时生效）：
    1. ``class_iris``（文档级）：非空时仅在该类子集内匹配（320→~50）；
    2. ``predict_segment_classes``（段级）：在文档级子集上排名，取 top-K 进一步缩窄（~50→~15）。
    ``class_iris`` 为 None 时跳过段级缩窄，保持全量匹配（320 类）。
    归类返回 ``None`` 的 span 被丢弃（无可信本体类）。
    """
    # 段级类预测：仅在文档级缩窄已生效时才进一步按段缩窄，
    # 避免 320→15 过度丢弃（原设计意图为 ~50→~15）。
    seg_top_classes: list[set[str] | None] | None = None
    if class_iris is not None:
        seg_top_classes = predict_segment_classes(
            segment_texts, engine, class_iris,
        )

    result: list[list[dict]] = [[] for _ in all_spans]
    for seg_idx, spans in enumerate(all_spans):
        if not spans:
            continue
        seg_text = segment_texts[seg_idx] if seg_idx < len(segment_texts) else ""
        seg_iris = class_iris
        if seg_top_classes is not None and seg_idx < len(seg_top_classes):
            seg_iris = seg_top_classes[seg_idx] or class_iris

        filtered: list[tuple[int, dict]] = [
            (si, span) for si, span in enumerate(spans)
            if not _is_non_entity_span(span["text"])
        ]
        if not filtered:
            continue

        # --- Pass-1: raw text, 高置信直接匹配 ---
        raw_texts = [span["text"] for _, span in filtered]
        pass1 = type_spans(
            raw_texts, engine,
            threshold=_DIRECT_MATCH_THRESHOLD, class_iris=seg_iris,
        )

        # --- Pass-2: 上下文辅助匹配（仅对 Pass-1 未命中的非短 span） ---
        need_pass2: list[int] = [
            fi for fi, match in enumerate(pass1)
            if match is None and not _is_short_span(filtered[fi][1]["text"])
        ]
        pass2_results: dict[int, dict] = {}
        if need_pass2:
            ctx_texts = [
                _span_with_context(
                    seg_text, filtered[fi][1]["start"], filtered[fi][1]["end"],
                )
                for fi in need_pass2
            ]
            typed2 = type_spans(ctx_texts, engine, class_iris=seg_iris)
            for j, fi in enumerate(need_pass2):
                if typed2[j] is not None:
                    pass2_results[fi] = typed2[j]

        # --- 合并 ---
        for fi, (_si, span) in enumerate(filtered):
            match = pass1[fi] or pass2_results.get(fi)
            if match is None:
                continue
            result[seg_idx].append({
                "start": span["start"],
                "end": span["end"],
                "text": span["text"],
                "label": match["label"],
                "className": match["label"],
                "score": match["score"],
                "iri": match["iri"],
            })
    return result


_PROPERTY_CONTEXT_WINDOW = 200


def _property_schema_for_class(engine, class_iri: str) -> dict:
    """从类 IRI 派生属性标签集（供 GLiNER 三元组抽取）。

    使用 ``get_data_properties_by_domain`` 规避 owlready2 ``get_class_properties`` bug。
    """
    props = engine.get_data_properties_by_domain(class_iri)
    labels: list[str] = []
    label_to_iri: dict[str, str] = {}
    for p in props:
        label = (p.get("label") or p.get("name") or "").strip()
        iri = p.get("iri")
        if label and iri and label not in label_to_iri:
            label_to_iri[label] = iri
            labels.append(label)
    return {"labels": labels, "label_to_iri": label_to_iri}


def _extract_property_triples(
    typed_spans: list[list[dict]],
    segment_texts: list[str],
    engine,
) -> list[dict]:
    """阶段三：对每个已归类实体 span，以其本体类数据属性为标签跑 GLiNER 抽取属性值。

    返回三元组列表，每个三元组含实体定位信息 + 抽取到的 ``(property_iri, value)`` 对。
    GLiNER 不可用 / 类无数据属性 → 该实体 properties 为空列表（结构保留、优雅降级）。
    """
    extractor = _get_extractor()
    if not extractor or not extractor.is_available():
        return []

    by_class: dict[str, list[tuple[int, dict]]] = {}
    for seg_idx, spans in enumerate(typed_spans):
        for span in spans:
            by_class.setdefault(span["iri"], []).append((seg_idx, span))

    triples: list[dict] = []
    for class_iri, entries in by_class.items():
        schema = _property_schema_for_class(engine, class_iri)
        if not schema["labels"]:
            for seg_idx, span in entries:
                triples.append({
                    "entity_text": span["text"],
                    "entity_class_iri": class_iri,
                    "entity_class_label": span["label"],
                    "segment_index": seg_idx,
                    "span_start": span["start"],
                    "span_end": span["end"],
                    "properties": [],
                })
            continue

        contexts: list[str] = []
        for seg_idx, span in entries:
            seg_text = segment_texts[seg_idx] if seg_idx < len(segment_texts) else ""
            contexts.append(_span_with_context(
                seg_text, span["start"], span["end"],
                window=_PROPERTY_CONTEXT_WINDOW,
            ))

        batch_spans = extractor.extract_batch_with_spans(
            contexts, schema["labels"], settings.gliner_threshold,
        )

        for (seg_idx, span), entity_spans in zip(entries, batch_spans):
            props: list[dict] = []
            seen_labels: set[str] = set()
            for es in entity_spans:
                label = es["label"]
                iri = schema["label_to_iri"].get(label)
                if iri and label not in seen_labels:
                    seen_labels.add(label)
                    props.append({"iri": iri, "label": label, "value": es["text"]})
            triples.append({
                "entity_text": span["text"],
                "entity_class_iri": class_iri,
                "entity_class_label": span["label"],
                "segment_index": seg_idx,
                "span_start": span["start"],
                "span_end": span["end"],
                "properties": props,
            })
    return triples


def _annotate_texts(
    texts: list[str],
    engine,
    progress_fn: Callable[[str], None] | None = None,
    should_pause_fn: Callable[[], bool] | None = None,
    checkpoint: dict | None = None,
    doc_class_iri: str | None = None,
) -> tuple[list[list[dict]], list[dict], dict | None]:
    """端到端三阶段：seed → GLiNER 定界 → 嵌入归类 → 属性三元组。

    ``doc_class_iri`` 非空时，Stage-2 候选类缩窄到文档类型子图（3 跳 ~100 类 vs 320）。
    Stage-1 始终使用通用种子标签——GLiNER 需要通用标签（"药物产品"、"设备"）才能广泛
    召回 span，而缩窄后的本体类名（"细胞毒API"）太专业、召回率极低。

    ``progress_fn`` 在每阶段开始时回调阶段名（"gliner"/"typing"/"triples"/"done"）。
    ``should_pause_fn`` 在每阶段间检查——返回 True 时中断并返回 checkpoint（第三元素）。
    ``checkpoint`` 传入已保存的中间结果以跳过已完成阶段。
    返回 ``(typed_spans, triples, checkpoint_or_None)``——checkpoint 为 None 表示正常完成。
    """
    from app.services.extraction.ontology_typer import relevant_classes_for_doc_type

    extractor = _get_extractor()
    # Stage-1 种子始终用通用标签（不随文档类型缩窄），保证 GLiNER 广泛召回
    labels = seed_labels(engine)
    class_iris = (
        relevant_classes_for_doc_type(engine, doc_class_iri) if doc_class_iri else None
    )

    # Stage 1: GLiNER span detection
    if checkpoint and checkpoint.get("completed_stage") in ("gliner", "typing"):
        raw = checkpoint["raw_spans"]
    else:
        if progress_fn:
            progress_fn("gliner")
        raw = _extract_spans_batch(extractor, texts, labels)
        if should_pause_fn and should_pause_fn():
            return [], [], {"completed_stage": "gliner", "raw_spans": raw}

    # Stage 2: embedding-based semantic typing (narrowed by doc type when available)
    if checkpoint and checkpoint.get("completed_stage") == "typing":
        typed = checkpoint["typed_spans"]
    else:
        if progress_fn:
            progress_fn("typing")
        typed = _type_and_filter_spans(raw, texts, engine, class_iris=class_iris)
        if should_pause_fn and should_pause_fn():
            return typed, [], {
                "completed_stage": "typing",
                "raw_spans": raw,
                "typed_spans": [[dict(s) for s in seg] for seg in typed],
            }

    # Stage 3: property triple extraction
    if progress_fn:
        progress_fn("triples")
    triples = _extract_property_triples(typed, texts, engine)
    if progress_fn:
        progress_fn("done")

    return typed, triples, None


# Word 段落样式名 → tiptap heading 层级（保留文档大纲）；中英文内置样式名都识别。
def _heading_level(style_name: str | None) -> int:
    if not style_name:
        return 0
    s = style_name.strip()
    low = s.lower()
    if low == "title" or s == "标题":
        return 1
    if low == "subtitle" or s == "副标题":
        return 2
    if "heading" in low or s.startswith("标题"):
        m = re.search(r"\d+", s)
        if m:
            return max(1, min(6, int(m.group())))
    return 0


_PT = 12700  # 1 point = 12700 EMU


def _font_size_heading_level(font_size_emu: int | None) -> int:
    """字号 → heading 层级启发式（段落样式名无法确定时的回退）。

    中文文档常用正文样式 + 手动字号排版；>=20pt 二号/小二→h1，>=16pt 三号→h2，>=14pt 四号→h3。
    """
    if not font_size_emu:
        return 0
    pt = font_size_emu / _PT
    if pt >= 20:
        return 1
    if pt >= 16:
        return 2
    if pt >= 14:
        return 3
    return 0


def _para_font_size(para) -> int | None:
    """段落有效字号（EMU）：首个非空 run → 段落样式字号。"""
    for r in para.runs:
        if (r.text or "").strip() and r.font.size:
            return r.font.size
    try:
        if para.style and para.style.font.size:
            return para.style.font.size
    except AttributeError:
        pass
    return None


# Word 段落对齐 → tiptap textAlign（左对齐为默认 → None 省略）。键为 WD_ALIGN_PARAGRAPH 枚举值。
_ALIGN_MAP = {1: "center", 2: "right", 3: "justify"}


def _para_align(para) -> str | None:
    a = para.alignment
    if a is None:
        return None
    try:
        return _ALIGN_MAP.get(int(a))
    except (TypeError, ValueError):
        return None


def _style_font_attr(para, attr: str):
    """段落样式字体属性（bold/italic/underline）；style 缺失或无属性 → None。"""
    try:
        return getattr(para.style.font, attr, None) if para.style else None
    except AttributeError:
        return None


def _para_runs_and_text(para) -> tuple[str, list[tuple[int, int, list[str]]]]:
    """由 run 自行拼接段落文本（而非 para.text），保证文本与 run 偏移严格对齐，
    使行内样式（粗体/斜体/下划线/删除线）能与 NER span 按字符边界无缝合并。

    run 属性为 None 时表示"继承段落样式"，需回退到 para.style.font 解析有效值。
    """
    style_bold = _style_font_attr(para, "bold")
    style_italic = _style_font_attr(para, "italic")
    style_underline = _style_font_attr(para, "underline")

    parts: list[str] = []
    runs: list[tuple[int, int, list[str]]] = []
    cursor = 0
    for r in para.runs:
        t = r.text or ""
        if not t:
            continue
        marks: list[str] = []
        bold = r.bold if r.bold is not None else style_bold
        if bold:
            marks.append("bold")
        italic = r.italic if r.italic is not None else style_italic
        if italic:
            marks.append("italic")
        underline = r.underline if r.underline is not None else style_underline
        if underline:
            marks.append("underline")
        if getattr(r.font, "strike", None):
            marks.append("strike")
        parts.append(t)
        runs.append((cursor, cursor + len(t), marks))
        cursor += len(t)
    return "".join(parts), runs


def _inline_nodes(
    text: str,
    spans: list[dict],
    runs: list[tuple[int, int, list[str]]] | None = None,
) -> list[dict]:
    """文本 + 行内样式 run + NER span → tiptap text 节点序列。

    按所有 run/span 边界切段，每段叠加其覆盖的行内样式 mark（bold/italic/…）与
    entity-annotation mark（落在某 span 内时），从而既保留 Word 行内格式又保留实体标注。
    """
    if not text:
        return []
    runs = runs or []
    bounds = {0, len(text)}
    for s, e, _ in runs:
        bounds.update((s, e))
    for sp in spans:
        bounds.update((sp["start"], sp["end"]))
    points = sorted(b for b in bounds if 0 <= b <= len(text))

    nodes: list[dict] = []
    for a, b in zip(points, points[1:]):
        if a >= b:
            continue
        marks: list[dict] = []
        for rs, re_, rmarks in runs:
            if rs <= a < re_:
                marks.extend({"type": m} for m in rmarks)
                break
        for sp in spans:
            if sp["start"] <= a < sp["end"]:
                marks.append({
                    "type": "entity-annotation",
                    "attrs": {
                        "label": sp["label"],
                        "className": sp.get("className", sp["label"]),
                        "score": sp["score"],
                    },
                })
                break
        node: dict = {"type": "text", "text": text[a:b]}
        if marks:
            node["marks"] = marks
        nodes.append(node)
    return nodes


def _text_to_tiptap_nodes(text: str, spans: list[dict]) -> list[dict]:
    """纯文本 + NER span → tiptap text 节点（无行内样式，表格单元格用）。"""
    return _inline_nodes(text, spans)


# ---------------------------------------------------------------------------
# Table processing helpers (009-word-table-ner-optimize)
# ---------------------------------------------------------------------------

def _detect_header_rows(table) -> int:
    """推断表格表头区域行数。

    无合并特征 → 1（仅第一行为表头）。首行含水平合并（gridSpan > 1）→ 继续扫描后续行，
    直至遇到无水平合并特征的行，返回总表头行数。
    """
    from docx.oxml.ns import qn

    rows = table.rows
    if not rows:
        return 0

    def _row_has_hmerge(row) -> bool:
        for tc_elem in row._tr.iterchildren(qn("w:tc")):
            tc_pr = tc_elem.find(qn("w:tcPr"))
            if tc_pr is not None:
                grid_span = tc_pr.find(qn("w:gridSpan"))
                if grid_span is not None:
                    val = grid_span.get(qn("w:val"))
                    if val is not None and int(val) > 1:
                        return True
        return False

    if not _row_has_hmerge(rows[0]):
        return 1

    header_count = 1
    for ri in range(1, len(rows)):
        if _row_has_hmerge(rows[ri]):
            header_count += 1
        else:
            header_count += 1
            break
    return header_count


def _find_table_caption(body_element, table_element) -> str | None:
    """返回表格前紧邻段落的文本（如"表 3：原料药规格"），作为语义上下文前缀。"""
    from docx.oxml.ns import qn

    prev = table_element.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        texts = [node.text or "" for node in prev.iter(qn("w:t"))]
        caption = "".join(texts).strip()
        if caption:
            return caption
    return None


def _is_vmerge_continue(tc_elem) -> bool:
    """检测纵向合并续行单元格（应跳过以避免重复 NER）。

    接受原始 ``<w:tc>`` XML 元素（而非 python-docx Cell），因为 python-docx 的
    ``Row.cells`` 对合并单元格返回 master cell，无法检测 continuation。
    """
    from docx.oxml.ns import qn

    tc_pr = tc_elem.find(qn("w:tcPr"))
    if tc_pr is None:
        return False
    vmerge = tc_pr.find(qn("w:vMerge"))
    if vmerge is None:
        return False
    val = vmerge.get(qn("w:val"))
    return val != "restart"


def _tc_text(tc_elem) -> str:
    """从原始 ``<w:tc>`` 元素提取文本内容（多段落用空格分隔，排除嵌套表格文本）。"""
    from docx.oxml.ns import qn

    paras: list[str] = []
    for p_elem in tc_elem.iterchildren(qn("w:p")):
        p_text = "".join(node.text or "" for node in p_elem.iter(qn("w:t")))
        if p_text.strip():
            paras.append(p_text.strip())
    return " ".join(paras)


def _build_row_segment(
    row, headers: list[str], caption: str | None = None,
    skip_vmerge: bool = False,
) -> tuple[str, list[tuple[int, int, int]]]:
    """将一行所有单元格拼接为行级文本段。

    遍历 ``row._tr`` 的原始 ``<w:tc>`` 子元素（而非 ``row.cells``），因为
    python-docx 的 Cell 访问对纵向合并单元格返回 master cell，导致续行无法检测。

    返回 ``(segment_text, cell_offsets)``，其中
    ``cell_offsets`` = ``[(cell_start, cell_end, col_idx), ...]``，
    记录每个 cell 内容在行级文本中的位置，用于 NER 后偏移校正。
    """
    from docx.oxml.ns import qn

    parts: list[str] = []
    cell_offsets: list[tuple[int, int, int]] = []
    cursor = 0

    for ci, tc_elem in enumerate(row._tr.iterchildren(qn("w:tc"))):
        if skip_vmerge and _is_vmerge_continue(tc_elem):
            continue
        cell_text = _tc_text(tc_elem)
        if not cell_text:
            continue
        hdr = headers[ci] if ci < len(headers) else ""
        if hdr:
            fragment = f"{hdr}：{cell_text}"
        else:
            fragment = cell_text

        if parts:
            sep = " | "
            cursor += len(sep)
            parts.append(sep)

        cell_offsets.append((cursor, cursor + len(fragment), ci))
        parts.append(fragment)
        cursor += len(fragment)

    segment = "".join(parts)
    if caption and segment:
        prefix = f"{caption} — "
        segment = prefix + segment
        cell_offsets = [
            (s + len(prefix), e + len(prefix), ci) for s, e, ci in cell_offsets
        ]

    return segment, cell_offsets


def _correct_span_offsets(
    spans: list[dict],
    cell_offsets: list[tuple[int, int, int]],
    headers: list[str],
) -> list[tuple[int, dict]]:
    """将行级 NER span 映射回 cell 内坐标。

    返回 ``[(col_idx, adjusted_span), ...]``，其中 adjusted_span 的
    start/end 相对于原始 cell 文本（不含表头前缀）。
    """
    results: list[tuple[int, dict]] = []
    for span in spans:
        for cell_start, cell_end, col_idx in cell_offsets:
            if span["start"] >= cell_start and span["end"] <= cell_end:
                hdr = headers[col_idx] if col_idx < len(headers) else ""
                prefix_len = len(f"{hdr}：") if hdr else 0
                ns = span["start"] - cell_start - prefix_len
                ne = span["end"] - cell_start - prefix_len
                if ns >= 0 and ne > ns:
                    results.append((col_idx, {**span, "start": ns, "end": ne}))
                break
    return results


def _collect_nested_segments(
    tbl_elem,
    caption: str | None,
    all_texts: list[str],
    nested_info: dict[tuple[int, int], list[dict]],
    parent_ri: int,
    parent_ci: int,
    depth: int = 0,
    max_depth: int = 5,
) -> None:
    """发现嵌套 ``<w:tbl>`` 并将其行级 segment 追加到 all_texts。"""
    from docx.oxml.ns import qn
    from docx.table import Table as DocxTable

    if depth >= max_depth:
        return
    try:
        inner_table = DocxTable(tbl_elem, None)
        rows = inner_table.rows
        if not rows:
            return
    except Exception:
        return

    n_hdr_count = _detect_header_rows(inner_table)
    n_headers: list[str] = []
    if n_hdr_count > 0 and len(rows) > 0:
        last_hdr = rows[min(n_hdr_count - 1, len(rows) - 1)]
        n_headers = [cell.text.strip() for cell in last_hdr.cells]

    n_seg_base = len(all_texts)
    n_row_seg_map: dict[int, int] = {}
    n_seg_cell_offsets: list[list[tuple[int, int, int]]] = []

    for nri, nrow in enumerate(rows):
        if nri < n_hdr_count:
            continue
        n_seg_text, n_cell_offs = _build_row_segment(
            nrow, n_headers, caption, skip_vmerge=True,
        )
        if not n_seg_text.strip():
            continue
        n_row_seg_map[nri] = len(all_texts) - n_seg_base
        all_texts.append(n_seg_text)
        n_seg_cell_offsets.append(n_cell_offs)

    if n_row_seg_map:
        tr_elems = [r._tr for r in rows]
        nested_info.setdefault((parent_ri, parent_ci), []).append({
            "seg_base": n_seg_base,
            "header_count": n_hdr_count,
            "headers": n_headers,
            "row_seg_map": n_row_seg_map,
            "seg_cell_offsets": n_seg_cell_offsets,
            "tr_elems": tr_elems,
        })

    for nri, nrow in enumerate(rows):
        if nri < n_hdr_count:
            continue
        for tc in nrow._tr.iterchildren(qn("w:tc")):
            for deeper_tbl in tc.findall(qn("w:tbl")):
                _collect_nested_segments(
                    deeper_tbl, caption, all_texts,
                    nested_info, parent_ri, parent_ci,
                    depth + 1, max_depth,
                )


def _render_nested_table(nested_info: dict, all_spans: list[list[dict]]) -> dict | None:
    """将嵌套表格渲染为 tiptap table 节点。"""
    from docx.oxml.ns import qn

    seg_base = nested_info["seg_base"]
    hdr_count = nested_info["header_count"]
    hdrs = nested_info["headers"]
    row_seg_map = nested_info["row_seg_map"]
    seg_offs = nested_info["seg_cell_offsets"]
    tr_elems = nested_info["tr_elems"]

    rows_data: list[dict] = []
    for tri, tr_elem in enumerate(tr_elems):
        cells: list[dict] = []
        tc_elems = list(tr_elem.iterchildren(qn("w:tc")))

        if tri < hdr_count:
            for tc in tc_elems:
                ct = _tc_text(tc)
                cchildren = _text_to_tiptap_nodes(ct, [])
                cells.append({
                    "type": "tableCell",
                    "content": [{"type": "paragraph", "content": cchildren}],
                })
        elif tri in row_seg_map:
            seg_offset = row_seg_map[tri]
            abs_idx = seg_base + seg_offset
            offsets = seg_offs[seg_offset]

            corrected = _correct_span_offsets(all_spans[abs_idx], offsets, hdrs)
            cell_span_map: dict[int, list[dict]] = {}
            for col_idx, adj_span in corrected:
                cell_span_map.setdefault(col_idx, []).append(adj_span)

            for ci, tc in enumerate(tc_elems):
                ct = "" if _is_vmerge_continue(tc) else _tc_text(tc)
                c_spans = cell_span_map.get(ci, [])
                cchildren = _text_to_tiptap_nodes(ct, c_spans)
                cells.append({
                    "type": "tableCell",
                    "content": [{"type": "paragraph", "content": cchildren}],
                })
        else:
            for tc in tc_elems:
                cells.append({
                    "type": "tableCell",
                    "content": [{"type": "paragraph"}],
                })

        if cells:
            rows_data.append({"type": "tableRow", "content": cells})

    return {"type": "table", "content": rows_data} if rows_data else None


def annotate_word(
    file_path: str | Path,
    engine,
    progress_fn: Callable[[str], None] | None = None,
    should_pause_fn: Callable[[], bool] | None = None,
    checkpoint: dict | None = None,
    doc_class_iri: str | None = None,
) -> tuple[dict, list[str], list[dict], dict | None]:
    """解析 Word 文档 → tiptap ProseMirror JSON，三阶段 NER 标注实体 + 属性三元组。

    ``doc_class_iri`` 非空时，NER 种子标签和候选类缩窄到文档类型的本体子图，
    大幅提升 Stage-2 嵌入归类精度（320 类 → ~50 类）。

    表格采用行级上下文拼接：同行单元格以 ``hdr：val | …`` 格式拼接为一个 segment
    送入 NER，NER 后将 span 坐标校正回各 cell 内坐标，前端 tiptap 渲染不受影响。
    表头行仅作列名前缀材料，不参与 NER。

    返回 ``(doc_json, warnings, triples, checkpoint_or_None)``。
    checkpoint 非 None 表示标注被暂停（doc_json 仍为完整结构，但标注可能不完整）。
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(file_path))

    _paras = {p._element: p for p in doc.paragraphs}
    _tables = {t._element: t for t in doc.tables}

    # Pass 1：按文档 body 顺序收集段落文本与表格行级 segment。
    elements: list[dict] = []
    all_texts: list[str] = []

    for child in doc.element.body:
        if child in _paras:
            para = _paras[child]
            text, runs = _para_runs_and_text(para)
            if text.strip():
                level = _heading_level(para.style.name if para.style else None)
                if not level:
                    level = _font_size_heading_level(_para_font_size(para))
                elements.append({
                    "kind": "para",
                    "level": level,
                    "align": _para_align(para),
                    "runs": runs,
                    "text_idx": len(all_texts),
                })
                all_texts.append(text)
            else:
                elements.append({"kind": "empty"})
        elif child in _tables:
            table = _tables[child]
            header_count = _detect_header_rows(table)
            caption = _find_table_caption(doc.element.body, table._element)

            headers: list[str] = []
            if header_count > 0 and len(table.rows) > 0:
                last_hdr = table.rows[min(header_count - 1, len(table.rows) - 1)]
                headers = [cell.text.strip() for cell in last_hdr.cells]

            seg_base = len(all_texts)
            row_seg_map: dict[int, int] = {}
            seg_cell_offsets: list[list[tuple[int, int, int]]] = []

            for ri, row in enumerate(table.rows):
                if ri < header_count:
                    continue
                seg_text, cell_offs = _build_row_segment(
                    row, headers, caption, skip_vmerge=True,
                )
                if not seg_text.strip():
                    continue
                row_seg_map[ri] = len(all_texts) - seg_base
                all_texts.append(seg_text)
                seg_cell_offsets.append(cell_offs)

            nested_tables_info: dict[tuple[int, int], list[dict]] = {}
            for ri, row in enumerate(table.rows):
                if ri < header_count:
                    continue
                for ci, tc_elem_n in enumerate(
                    row._tr.iterchildren(qn("w:tc"))
                ):
                    for itbl in tc_elem_n.findall(qn("w:tbl")):
                        _collect_nested_segments(
                            itbl, caption, all_texts,
                            nested_tables_info, ri, ci,
                        )

            elements.append({
                "kind": "table",
                "table_ref": table,
                "seg_base": seg_base,
                "header_count": header_count,
                "headers": headers,
                "row_seg_map": row_seg_map,
                "seg_cell_offsets": seg_cell_offsets,
                "nested_tables": nested_tables_info,
            })

    # 端到端三阶段标注（一次 batch）——行级 segment 使 _span_with_context
    # 的 window=40 覆盖同行相邻列上下文（research R5）。
    all_spans, triples, ckpt = _annotate_texts(
        all_texts, engine, progress_fn, should_pause_fn, checkpoint,
        doc_class_iri=doc_class_iri,
    )

    # Pass 2：按记录顺序组装 tiptap 节点。
    content: list[dict] = []
    for elem in elements:
        if elem["kind"] == "empty":
            content.append({"type": "paragraph"})
            continue
        if elem["kind"] == "para":
            idx = elem["text_idx"]
            children = _inline_nodes(all_texts[idx], all_spans[idx], elem["runs"])
            level = elem["level"]
            node: dict = (
                {"type": "heading", "attrs": {"level": level}}
                if level
                else {"type": "paragraph"}
            )
            if elem["align"]:
                node.setdefault("attrs", {})["textAlign"] = elem["align"]
            node["content"] = children
            content.append(node)
        elif elem["kind"] == "table":
            table_ref = elem["table_ref"]
            seg_base = elem["seg_base"]
            hdr_count = elem["header_count"]
            hdrs = elem["headers"]
            row_seg_map = elem["row_seg_map"]
            seg_offs = elem["seg_cell_offsets"]
            nested_tables = elem.get("nested_tables", {})

            rows_data: list[dict] = []
            for ri, row in enumerate(table_ref.rows):
                cells: list[dict] = []
                tc_elems = list(row._tr.iterchildren(qn("w:tc")))
                if ri < hdr_count:
                    for tc in tc_elems:
                        ct = _tc_text(tc)
                        cchildren = _text_to_tiptap_nodes(ct, [])
                        cells.append({
                            "type": "tableCell",
                            "content": [{"type": "paragraph", "content": cchildren}],
                        })
                elif ri in row_seg_map:
                    seg_offset = row_seg_map[ri]
                    abs_idx = seg_base + seg_offset
                    offsets = seg_offs[seg_offset]

                    corrected = _correct_span_offsets(
                        all_spans[abs_idx], offsets, hdrs,
                    )
                    cell_span_map: dict[int, list[dict]] = {}
                    for col_idx, adj_span in corrected:
                        cell_span_map.setdefault(col_idx, []).append(adj_span)

                    for ci, tc in enumerate(tc_elems):
                        ct = "" if _is_vmerge_continue(tc) else _tc_text(tc)
                        c_spans = cell_span_map.get(ci, [])
                        cchildren = _text_to_tiptap_nodes(ct, c_spans)
                        cell_content: list[dict] = [
                            {"type": "paragraph", "content": cchildren},
                        ]
                        for n_info in nested_tables.get((ri, ci), []):
                            n_node = _render_nested_table(n_info, all_spans)
                            if n_node:
                                cell_content.append(n_node)
                        cells.append({
                            "type": "tableCell",
                            "content": cell_content,
                        })
                else:
                    for cell in row.cells:
                        cells.append({
                            "type": "tableCell",
                            "content": [{"type": "paragraph"}],
                        })

                rows_data.append({"type": "tableRow", "content": cells})

            if rows_data:
                content.append({"type": "table", "content": rows_data})

    return {"type": "doc", "content": content}, _warnings(all_texts), triples, ckpt


def annotate_excel(
    file_path: str | Path,
    engine,
    ner_columns: list[str] | None = None,
    progress_fn: Callable[[str], None] | None = None,
    should_pause_fn: Callable[[], bool] | None = None,
    checkpoint: dict | None = None,
) -> tuple[dict[str, Any], list[str], list[dict], dict | None]:
    """解析 Excel → 结构化行数据 + 三阶段 NER 标注 + 属性三元组。

    返回 ``(data, warnings, triples, checkpoint_or_None)``。
    """
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        return {"headers": [], "rows": []}, [], [], None

    rows_iter = ws.iter_rows(values_only=False)

    header_row = next(rows_iter, None)
    if header_row is None:
        wb.close()
        return {"headers": [], "rows": []}, [], [], None

    headers = [str(c.value or f"col_{i}") for i, c in enumerate(header_row)]
    ner_set = set(ner_columns) if ner_columns else set(headers)

    # Pass 1：收集所有行/列单元格文本，记录哪些需要 NER。
    raw_rows: list[dict[str, str]] = []
    ner_texts: list[str] = []
    ner_slots: list[tuple[int, str]] = []
    for row in rows_iter:
        cells: dict[str, str] = {}
        for i, cell in enumerate(row):
            if i >= len(headers):
                break
            col_name = headers[i]
            value = cell.value
            text = str(value) if value is not None else ""
            cells[col_name] = text
            if col_name in ner_set and text:
                ner_slots.append((len(raw_rows), col_name))
                ner_texts.append(text)
        raw_rows.append(cells)
    wb.close()

    # 端到端三阶段标注所有需标注单元格。
    ner_spans, triples, ckpt = _annotate_texts(
        ner_texts, engine, progress_fn, should_pause_fn, checkpoint,
    )
    spans_by_slot: dict[tuple[int, str], list[dict]] = {
        slot: spans for slot, spans in zip(ner_slots, ner_spans)
    }

    # Pass 2：组装行数据。
    rows: list[dict] = []
    for r_idx, cells in enumerate(raw_rows):
        out_cells: dict[str, Any] = {}
        for col_name, text in cells.items():
            annotations = spans_by_slot.get((r_idx, col_name), [])
            out_cells[col_name] = {"value": text, "annotations": annotations}
        rows.append(out_cells)

    return {"headers": headers, "rows": rows}, _warnings(ner_texts), triples, ckpt


def _warnings(texts: list[str]) -> list[str]:
    """记录 get_class_properties domain bug（仅当确有文本走了 NER 时上浮，避免空噪声）。"""
    return [GET_CLASS_PROPERTIES_DOMAIN_BUG] if any(texts) else []
