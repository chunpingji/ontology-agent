"""Word 文档结构薄解析 —— 供文档级分类与规则式关系抽取共用。

三阶段 NER（``document_annotator``）把 Word 渲染成 tiptap JSON 用于高亮预览；而文档
级分类与关系抽取需要的是另一种视角：**标题层级 + 章节归属 + 表格的表头签名/行**。
本模块产出这样一个轻量结构视图（不重复 tiptap 构建逻辑），复用
``document_annotator._heading_level`` 做样式名→层级映射，保持与预览同口径的大纲。

产出 ``DocStructure``：

- ``title``：首个一级标题文本（无则文件名 stem）。
- ``sections``：``[{heading, level, paras}]``——按标题切分，每节含其下正文段落文本。
- ``tables``：``[{headers, rows, cells, ncols}]``——``headers`` 为首行单元格文本，
  ``rows`` 为「表头→单元格」字典列表（表头匹配用），``cells`` 为原始二维文本。
- ``paragraphs``：全部非空段落文本（扁平，供分类 haystack）。
- ``headings``：全部标题文本（扁平，供分类 haystack）。

``python-docx`` 的 ``doc.paragraphs`` 保序、``doc.tables`` 独立——本模块据此从段落流
切章节、从表格流取表头签名；关系抽取按**表头签名**（而非脆弱的「表N」题注）定位表格。
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from app.services.extraction.document_annotator import _heading_level


@dataclass
class DocSection:
    heading: str
    level: int
    paras: list[str] = field(default_factory=list)


@dataclass
class DocTable:
    headers: list[str]
    rows: list[dict[str, str]]
    cells: list[list[str]]

    @property
    def ncols(self) -> int:
        return len(self.headers)

    @property
    def header_sig(self) -> str:
        """表头签名：所有表头单元格拼接，用于子串匹配定位目标表。"""
        return " | ".join(self.headers)

    def has_headers(self, *needles: str) -> bool:
        """表头签名是否同时包含全部 ``needles``（表格定位）。"""
        sig = self.header_sig
        return all(n in sig for n in needles)


@dataclass
class DocStructure:
    title: str
    sections: list[DocSection]
    tables: list[DocTable]
    paragraphs: list[str]
    headings: list[str]

    def find_section(self, *needles: str) -> DocSection | None:
        """返回首个标题包含任一 ``needles`` 的章节。"""
        for sec in self.sections:
            if any(n in sec.heading for n in needles):
                return sec
        return None

    def find_table(self, *needles: str) -> DocTable | None:
        """返回首个表头同时包含全部 ``needles`` 的表格（表头签名匹配）。"""
        for tbl in self.tables:
            if tbl.has_headers(*needles):
                return tbl
        return None


def _cell_text(cell) -> str:
    """单元格文本：合并多段、去首尾空白（合并单元格会重复同一 cell 对象，无妨）。"""
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def _table_to_struct(table) -> DocTable:
    cells: list[list[str]] = [[_cell_text(c) for c in row.cells] for row in table.rows]
    headers = cells[0] if cells else []
    rows: list[dict[str, str]] = []
    for raw in cells[1:]:
        row = {}
        for i, val in enumerate(raw):
            key = headers[i] if i < len(headers) and headers[i] else f"col{i}"
            # 同名表头不覆盖：保留首个非空（罕见，防御性）。
            if key not in row or not row[key]:
                row[key] = val
        rows.append(row)
    return DocTable(headers=headers, rows=rows, cells=cells)


def parse_docx_structure(file_path: str | Path) -> DocStructure:
    """解析 .docx → ``DocStructure``（标题/章节/表格签名）。绝不抛出：失败返回空结构。"""
    path = Path(file_path)
    try:
        from docx import Document
    except Exception:  # pragma: no cover - python-docx 未安装的防御路径
        return DocStructure(path.stem, [], [], [], [])

    try:
        doc = Document(str(path))
    except Exception:  # pragma: no cover - 损坏/非 docx 的防御路径
        return DocStructure(path.stem, [], [], [], [])

    sections: list[DocSection] = []
    paragraphs: list[str] = []
    headings: list[str] = []
    title = ""

    current = DocSection(heading="", level=0)  # 首个标题前的前言段
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = para.style.name if para.style else None
        level = _heading_level(style_name)
        paragraphs.append(text)
        if level > 0:
            headings.append(text)
            if level == 1 and not title:
                title = text
            if current.heading or current.paras:
                sections.append(current)
            current = DocSection(heading=text, level=level)
        else:
            current.paras.append(text)
    if current.heading or current.paras:
        sections.append(current)

    tables = [_table_to_struct(t) for t in doc.tables]

    return DocStructure(
        title=title or path.stem,
        sections=sections,
        tables=tables,
        paragraphs=paragraphs,
        headings=headings,
    )
