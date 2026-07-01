"""Render a RiskReport to .docx bytes (010, FR-005).

Produces a QS-A-020F05 formatted Word document with:
- Header with document number/revision/date + coverage summary (AST-5)
- SECTION I: Subject description, equipment tables by workshop, assessment table
- SECTION II: Placeholders for risk review and conclusion

When a :class:`CoverageManifest` is supplied (AST-5), the document carries the
no-omission evidence visibly: a coverage banner, an outstanding-materials list,
and red highlighting on any "⚠ 待评估" cell — so a missing material can never be
silently absent from the rendered report.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.reporting.coverage_validator import CoverageManifest
from app.services.reporting.risk_report_generator import RiskReport

# Visual flag for missing/pending material (AST-5). Matches the warning glyph
# used by the generator's PENDING_LEVEL and the template's missing_placeholder.
_WARN_GLYPH = "⚠"
_WARN_COLOR = RGBColor(0xC0, 0x00, 0x00)


def render_risk_report(
    report: RiskReport, manifest: CoverageManifest | None = None
) -> bytes:
    """Render ``RiskReport`` dataclass to .docx bytes.

    ``manifest`` (AST-5) is optional for backward compatibility; when given, its
    coverage status is surfaced in the rendered document.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)

    _add_page_header(doc, report, manifest)
    _add_header(doc, report)
    _add_coverage_banner(doc, manifest)
    _add_section_one(doc, report)
    _add_assessment_table(doc, report)
    _add_outstanding_materials(doc, manifest)
    _add_section_two(doc, report)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_page_header(
    doc: Document, report: RiskReport, manifest: CoverageManifest | None = None
) -> None:
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{report.doc_no}  Rev.{report.revision}")
    run.font.size = Pt(8)
    run.font.name = "宋体"
    if report.effective_date:
        run2 = p.add_run(f"  |  {report.effective_date}")
        run2.font.size = Pt(8)
        run2.font.name = "宋体"
    if manifest is not None and manifest.has_omissions:
        warn = p.add_run(f"  |  {_WARN_GLYPH} {manifest.missing_required} 项待补充")
        warn.font.size = Pt(8)
        warn.font.name = "宋体"
        warn.bold = True
        warn.font.color.rgb = _WARN_COLOR


def _add_coverage_banner(doc: Document, manifest: CoverageManifest | None) -> None:
    """A one-line material-coverage summary directly under the title (AST-5)."""
    if manifest is None:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = [
        f"素材覆盖：共 {manifest.total_slots} 项",
        f"已填充 {manifest.filled}",
        f"推理 {manifest.inferred}",
        f"人工 {manifest.manual}",
        f"待补充 {manifest.missing_required}",
    ]
    if manifest.dismissed > 0:
        parts.append(f"不适用 {manifest.dismissed}")
    summary = " · ".join(parts)
    run = p.add_run(summary)
    run.font.size = Pt(9)
    run.font.name = "宋体"
    if manifest.has_omissions:
        run.bold = True
        run.font.color.rgb = _WARN_COLOR


def _add_header(doc: Document, report: RiskReport) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"风险评估表 ({report.doc_no})")
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"修订号: {report.revision}    生效日期: {report.effective_date}")


def _add_section_one(doc: Document, report: RiskReport) -> None:
    doc.add_heading("SECTION I  风险评估", level=2)

    doc.add_heading("1. 风险评估对象 Subject Description", level=3)
    doc.add_paragraph(report.subject_description or "（待补充）")

    doc.add_heading("2. 评估小组 Assessment Team", level=3)
    if report.team_members:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        hdr[0].text = "姓名\nName"
        hdr[1].text = "职务\nTitle"
        hdr[2].text = "签名\nSignature"
        for m in report.team_members:
            row = t.add_row().cells
            row[0].text = m.get("name", "")
            row[1].text = m.get("title", "")
    else:
        doc.add_paragraph("（待补充）")

    if report.equipment_tables:
        doc.add_heading("3. 设备一览表 Equipment List", level=3)
        for workshop, entries in report.equipment_tables.items():
            doc.add_paragraph(f"● {workshop}", style="List Bullet")
            t = doc.add_table(rows=1, cols=5)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = t.rows[0].cells
            hdr[0].text = "序号\nNo."
            hdr[1].text = "设备编号\nEquipment ID"
            hdr[2].text = "设备名称\nName"
            hdr[3].text = "规格\nSpecification"
            hdr[4].text = "材质\nMaterial"
            for e in entries:
                row = t.add_row().cells
                row[0].text = str(e.seq)
                row[1].text = e.equipment_id
                row[2].text = e.name
                row[3].text = e.spec
                row[4].text = e.material

        for note in report.equipment_notes:
            doc.add_paragraph(f"注: {note}", style="List Bullet")


def _add_assessment_table(doc: Document, report: RiskReport) -> None:
    doc.add_heading("4. 风险评估 Risk Assessment", level=3)

    headers = [
        "HazID\n风险类型",
        "Contributing Factors\n风险因素",
        "Pre-Control\n控制前风险等级",
        "Post-Control\n控制后风险等级",
        "Control Measures\n风险控制措施",
        "Traceability\n控制可追溯性",
        "Status\n风险状态",
    ]

    col_widths = [Cm(2.5), Cm(5.5), Cm(2.5), Cm(2.5), Cm(6.0), Cm(4.0), Cm(2.5)]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, w in enumerate(col_widths):
        t.columns[i].width = w
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].width = col_widths[i]
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "宋体"

    for row_data in report.assessment_rows:
        row = t.add_row().cells
        vals = [
            row_data.hazid, row_data.contributing_factors,
            row_data.pre_control_level, row_data.post_control_level,
            row_data.control_measures, row_data.traceability, row_data.status,
        ]
        for i, val in enumerate(vals):
            row[i].text = val
            pending = bool(val) and val.startswith(_WARN_GLYPH)
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "宋体"
                    if pending:  # AST-5: 待评估 cells flagged red+bold
                        run.bold = True
                        run.font.color.rgb = _WARN_COLOR

    doc.add_paragraph()
    doc.add_paragraph(f"QA 评论: {report.qa_comments or '（无）'}")


def _add_outstanding_materials(
    doc: Document, manifest: CoverageManifest | None
) -> None:
    """List every required slot that could not be filled (AST-5, no-omission proof).

    This subsection makes the G1 omissions explicit and reviewable: each missing
    required material is named with its dimension/source so a human knows exactly
    what to supply — the report never silently drops a required input.
    Dismissed slots are listed separately as "N/A（不适用）" (011 FR-API-006).
    """
    has_missing = manifest is not None and manifest.has_omissions
    has_dismissed = manifest is not None and manifest.dismissed > 0
    if not has_missing and not has_dismissed:
        return
    doc.add_heading("5. 待补充素材清单 Outstanding Materials", level=3)
    if has_missing:
        intro = doc.add_paragraph()
        run = intro.add_run(
            f"{_WARN_GLYPH} 以下 {manifest.missing_required} 项必填素材未能从抽取数据中确定，"
            "需人工补充后方可定论；在此之前相关结论标记为「待评估」。"
        )
        run.bold = True
        run.font.color.rgb = _WARN_COLOR
        for slot in manifest.missing_required_slots:
            label = slot.label or slot.slot_id
            text = f"{label}（{slot.slot_id}）"
            if slot.note:
                text += f" — {slot.note}"
            doc.add_paragraph(text, style="List Bullet")
    if has_dismissed:
        intro_d = doc.add_paragraph()
        run_d = intro_d.add_run(
            f"以下 {manifest.dismissed} 项已确认为不适用（N/A），不计入缺失。"
        )
        run_d.font.size = Pt(9)
        run_d.font.name = "宋体"
        for slot in manifest.dismissed_slots:
            label = slot.label or slot.slot_id
            doc.add_paragraph(f"{label}（{slot.slot_id}）— N/A（不适用）", style="List Bullet")


def _add_section_two(doc: Document, report: RiskReport) -> None:
    doc.add_heading("SECTION II  风险回顾 Risk Review", level=2)
    doc.add_paragraph(report.risk_review or "（定期回顾时补充）")

    doc.add_heading("结论 Conclusion", level=3)
    doc.add_paragraph(report.conclusion or "（待补充）")

    doc.add_heading("审批 Approvals", level=3)
    if report.approvers:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text = "角色\nRole"
        hdr[1].text = "姓名\nName"
        hdr[2].text = "签名\nSignature"
        hdr[3].text = "日期\nDate"
        for a in report.approvers:
            row = t.add_row().cells
            row[0].text = a.get("role", "")
            row[1].text = a.get("name", "")
    else:
        doc.add_paragraph("（待补充）")
