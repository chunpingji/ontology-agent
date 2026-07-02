"""013: parse-sample + suggest-slots endpoint contract tests.

- parse-sample: 确定性离线解析（角色门控 senior_analyst，但**不**受 llm 门控——
  LLM 关闭时也能预览结构）；
- suggest-slots: 三选一输入校验、flag/client 503、角色 403，以及 sample_content_json
  分支走服务端 tiptap→text + 结构锚点派生（mock 本地 LLM，零云端调用）。

Mock 约定与 test_slot_suggester 一致：在 OpenAI 客户端层 mock
``client.chat.completions.create``，跑真实的 suggest_slots → chat_with_schema。
"""

from __future__ import annotations

import json
from unittest.mock import MagicMock

import docx
import pytest
from pydantic import ValidationError

from app.schemas.extraction import SuggestSlotsRequest


# ── canned LLM rounds + mock client ──────────────────────────────────────────

_R1_OK = {
    "document_summary": "GMP 风险评估报告",
    "sections": [
        {
            "title": "评估对象",
            "groups": [
                {
                    "title": "药品信息",
                    "candidates": [{"label": "药品名称", "evidence_span": "XX注射液"}],
                },
            ],
        },
    ],
}

_R2_OK = {
    "slots": [
        {
            "slot_id": "drug_name",
            "label": "药品名称",
            "section": "评估对象",
            "group": "药品信息",
            "confidence": 0.92,
            "evidence_span": "XX注射液",
            "reason": "文档首段",
        },
    ],
    "skipped_duplicates": 0,
}


def _make_llm_client(round1: dict, round2: dict):
    """Mock OpenAI client: canned JSON for the two chat.completions.create calls."""
    client = MagicMock()
    responses = [round1, round2]
    n = {"i": 0}

    def _create(**kwargs):
        idx = min(n["i"], len(responses) - 1)
        n["i"] += 1
        resp = MagicMock()
        choice = MagicMock()
        choice.message.content = json.dumps(responses[idx], ensure_ascii=False)
        resp.choices = [choice]
        return resp

    client.chat.completions.create = _create
    return client


# 忠于结构的样例：标题 + 含证据片段的段落 → source_ref 应派生为 "§ 评估对象"。
_SAMPLE_DOC = {
    "type": "doc",
    "content": [
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "评估对象"}],
        },
        {
            "type": "paragraph",
            "content": [{"type": "text", "text": "本品为XX注射液，剂型为注射剂。"}],
        },
    ],
}

_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _docx_bytes(tmp_path) -> bytes:
    path = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_heading("评估对象", level=1)
    d.add_paragraph("本品为XX注射液，剂型为注射剂。")
    d.save(str(path))
    return path.read_bytes()


def _enable_llm(monkeypatch, client):
    """Flag on + local LLM returns the given mock client (endpoint imports at call time)."""
    from app.config import settings

    monkeypatch.setattr(settings, "llm_suggest_slots_enabled", True)
    monkeypatch.setattr(
        "app.services.llm.local_client.get_local_llm", lambda: client
    )


# ── SuggestSlotsRequest.model_post_init — 3-way "exactly one" (pure schema) ──

class TestSuggestSlotsRequestValidation:
    def test_zero_sources_rejected(self):
        with pytest.raises(ValidationError):
            SuggestSlotsRequest()

    def test_two_sources_rejected(self):
        with pytest.raises(ValidationError):
            SuggestSlotsRequest(document_text="x", sample_content_json={"type": "doc"})

    def test_exactly_one_ok(self):
        assert SuggestSlotsRequest(sample_content_json=_SAMPLE_DOC).sample_content_json
        assert SuggestSlotsRequest(document_text="hi").document_text == "hi"
        jid = "00000000-0000-0000-0000-000000000001"
        assert str(SuggestSlotsRequest(job_id=jid).job_id) == jid


# ── POST /api/ast-templates/parse-sample (role-gated, NOT flag-gated) ────────

class TestParseSampleEndpoint:
    def test_happy_path_returns_content_and_text(self, client, analyst_headers, tmp_path):
        resp = client.post(
            "/api/ast-templates/parse-sample",
            headers=analyst_headers,
            files={"file": ("sample.docx", _docx_bytes(tmp_path), _DOCX_MIME)},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["content_json"]["type"] == "doc"
        assert "评估对象" in body["plain_text"]

    def test_non_docx_rejected(self, client, analyst_headers):
        resp = client.post(
            "/api/ast-templates/parse-sample",
            headers=analyst_headers,
            files={"file": ("notes.txt", b"hello", "text/plain")},
        )
        assert resp.status_code == 422

    def test_role_gated(self, client, operator_headers, tmp_path):
        # 发送合法 docx，确保 403 来自角色门控而非缺参/文件类型。
        resp = client.post(
            "/api/ast-templates/parse-sample",
            headers=operator_headers,
            files={"file": ("sample.docx", _docx_bytes(tmp_path), _DOCX_MIME)},
        )
        assert resp.status_code == 403


# ── POST /api/ast-templates/suggest-slots ────────────────────────────────────

class TestSuggestSlotsEndpoint:
    def test_sample_content_json_path_binds_source_ref(
        self, client, analyst_headers, monkeypatch
    ):
        _enable_llm(monkeypatch, _make_llm_client(_R1_OK, _R2_OK))
        resp = client.post(
            "/api/ast-templates/suggest-slots",
            headers=analyst_headers,
            json={"sample_content_json": _SAMPLE_DOC},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["total_suggested"] == 1
        slots = [
            s for sec in body["sections"] for g in sec["groups"] for s in g["slots"]
        ]
        assert slots[0]["source_ref"] == "§ 评估对象"

    def test_flag_off_returns_503(self, client, analyst_headers, monkeypatch):
        from app.config import settings

        monkeypatch.setattr(settings, "llm_suggest_slots_enabled", False)
        resp = client.post(
            "/api/ast-templates/suggest-slots",
            headers=analyst_headers,
            json={"document_text": "some text"},
        )
        assert resp.status_code == 503

    def test_no_local_client_returns_503(self, client, analyst_headers, monkeypatch):
        from app.config import settings

        monkeypatch.setattr(settings, "llm_suggest_slots_enabled", True)
        monkeypatch.setattr(
            "app.services.llm.local_client.get_local_llm", lambda: None
        )
        resp = client.post(
            "/api/ast-templates/suggest-slots",
            headers=analyst_headers,
            json={"document_text": "some text"},
        )
        assert resp.status_code == 503

    def test_role_gated_403(self, client, operator_headers):
        # 角色依赖先于 flag/client 检查触发——无需 mock LLM。
        resp = client.post(
            "/api/ast-templates/suggest-slots",
            headers=operator_headers,
            json={"document_text": "some text"},
        )
        assert resp.status_code == 403

    def test_two_sources_returns_422(self, client, analyst_headers):
        resp = client.post(
            "/api/ast-templates/suggest-slots",
            headers=analyst_headers,
            json={"document_text": "x", "sample_content_json": {"type": "doc"}},
        )
        assert resp.status_code == 422
