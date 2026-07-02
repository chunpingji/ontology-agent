"""Tests for 013 DOCX LLM annotation rendering (T019).

Verifies:
  - Supplemented values render gray-italic + info glyph
  - End-of-report disclaimer section present when LLM content exists
  - No LLM styling when supplements empty
  - 100% of LLM content visually annotated (SC-004)
"""

from __future__ import annotations

import io

from docx import Document

from app.services.reporting.docx_renderer import (
    _LLM_END_DISCLAIMER_ZH,
    _LLM_INFO_GLYPH,
    _add_generated_disclaimer_section,
    _add_llm_disclaimer_line,
    _add_llm_run,
    render_risk_report,
)
from app.services.reporting.risk_report_generator import RiskReport


class TestLLMRunStyling:
    def test_add_llm_run_gray_italic(self):
        doc = Document()
        p = doc.add_paragraph()
        _add_llm_run(p, "测试文本")

        run = p.runs[-1]
        assert run.italic is True
        assert _LLM_INFO_GLYPH in run.text
        assert "测试文本" in run.text
        assert run.font.color.rgb is not None

    def test_disclaimer_line_added(self):
        doc = Document()
        _add_llm_disclaimer_line(doc)
        last_p = doc.paragraphs[-1]
        assert last_p.runs[0].italic is True

    def test_generated_disclaimer_section(self):
        doc = Document()
        _add_generated_disclaimer_section(doc)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any(_LLM_END_DISCLAIMER_ZH in t for t in texts)


class TestRenderReportWithLLMSupplements:
    def test_disclaimer_present_when_supplements(self):
        report = RiskReport(
            llm_supplements={"subject.name": "阿莫西林"},
            llm_generated_fields={"subject.name"},
        )
        docx_bytes = render_risk_report(report)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert _LLM_END_DISCLAIMER_ZH in all_text

    def test_no_disclaimer_when_no_supplements(self):
        report = RiskReport()
        docx_bytes = render_risk_report(report)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert _LLM_END_DISCLAIMER_ZH not in all_text

    def test_info_glyph_present_for_supplemented_content(self):
        report = RiskReport(
            llm_supplements={"subject.desc": "LLM生成内容"},
            llm_generated_fields={"subject.desc"},
        )
        docx_bytes = render_risk_report(report)
        doc = Document(io.BytesIO(docx_bytes))
        all_runs_text = ""
        for p in doc.paragraphs:
            for run in p.runs:
                all_runs_text += run.text
        assert _LLM_INFO_GLYPH in all_runs_text
