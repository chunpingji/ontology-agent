"""集成测试：Word 正文 prose 实体经本地 NER → instance 候选入审核队列（US2）。

覆盖 data-model §3.3 / FR-005/006/010/011：Word `paragraph` 段落经本地零样本 NER 召回为
`candidate_kind="instance"` 候选——`source_ref` 含 `#para` 溯源、`review_status="pending"`
不自动断言、携 `align_entity` 对齐结果/分数/分组键；与 `parse_action_from_text` 的 Action
候选**并存**；同标签多命中聚为多值 list。

以确定性 GLiNER 桩（替换 `get_gliner_extractor`）+ 假本体引擎（`get_class_detail` 派生
标签、`get_individuals` 返空）注入，不下载真实权重——对齐 test_schema_derivation.py /
test_doc_repo_pipeline.py 的注入式做法。
"""

from __future__ import annotations

import asyncio
from types import SimpleNamespace

import docx

from app.config import settings
from app.models.extraction import ExtractionCandidate, ExtractionConfig, ExtractionJob
from app.services.extraction import pipeline as pipeline_module
from app.services.extraction.pipeline import run_extraction_pipeline

DRUG = "https://ontology.pharma-gmp.cn/slpra/drug/DrugProduct"
IRI_ACTIVE = f"{DRUG}#activeIngredient"
IRI_DOSAGE = f"{DRUG}#dosageForm"
IRI_SPEC = f"{DRUG}#specification"

# 业务实体段、同标签多值段、条件式 Action 段（NER 对其无召回）。
PARA_ENTITY = "本品活性成分为化合物A，剂型为片剂，规格10mg。"
PARA_MULTI = "复方制剂含化合物X与化合物Y两种活性成分。"
PARA_ACTION = "若产品为高致敏药品，则必须执行专用化生产。"

# 桩 NER：按段落文本预置 {label: value|[values]}（真实 GLiNER 的零样本输出形态）。
_NER_BY_TEXT = {
    PARA_ENTITY: {"活性成分": "化合物A", "剂型": "片剂", "规格": "10mg"},
    PARA_MULTI: {"活性成分": ["化合物X", "化合物Y"]},
    PARA_ACTION: {},
}


class _FakeExtractor:
    """确定性 GLiNER 桩：按文本返回预置实体，仅保留 labels 中声明的键（标签驱动）。"""

    def is_available(self) -> bool:
        return True

    def extract_text(self, text, labels, threshold=None):
        result = _NER_BY_TEXT.get(text, {})
        return {k: v for k, v in result.items() if k in labels}


class _ProseEngine:
    """假本体引擎：`get_class_detail` 派生药品标签集；`get_individuals` 返空（对齐→new）。"""

    def get_class_detail(self, iri):
        return SimpleNamespace(data_properties=[
            {"iri": IRI_ACTIVE, "name": "activeIngredient", "label": "活性成分", "range": ["string"]},
            {"iri": IRI_DOSAGE, "name": "dosageForm", "label": "剂型", "range": ["string"]},
            {"iri": IRI_SPEC, "name": "specification", "label": "规格", "range": ["string"]},
        ])

    def get_individuals(self, iri):
        return []

    def __getattr__(self, name):  # 容忍其余引擎调用（align_entity 仅用上面两个）
        def _noop(*a, **k):
            return None
        return _noop


def _docx(path):
    doc = docx.Document()
    doc.add_paragraph(PARA_ENTITY)
    doc.add_paragraph(PARA_MULTI)
    doc.add_paragraph(PARA_ACTION)
    doc.save(path)
    return path


def _config(db):
    cfg = ExtractionConfig(
        name="内部药物实体", target_class_iri=DRUG, source_type="word",
        column_mapping={"活性成分": "activeIngredient", "剂型": "dosageForm",
                        "规格": "specification"},
    )
    db.add(cfg)
    db.commit()
    db.refresh(cfg)
    return cfg


def _job(db):
    job = ExtractionJob(source_type="word", source_filename="SOP.docx",
                        source_config={}, status="pending")
    db.add(job)
    db.commit()
    db.refresh(job)
    return job


def _run(db, monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "llm_cloud_enabled", False)
    monkeypatch.setattr(settings, "anthropic_api_key", "")
    monkeypatch.setattr(pipeline_module, "get_gliner_extractor", lambda: _FakeExtractor())
    cfg, job = _config(db), _job(db)
    asyncio.run(run_extraction_pipeline(job, cfg, _docx(tmp_path / "SOP.docx"),
                                        _ProseEngine(), db))
    db.refresh(job)
    return job


def _candidates(db, job):
    return db.query(ExtractionCandidate).filter(ExtractionCandidate.job_id == job.id).all()


def test_prose_entities_become_pending_instance_candidates(db, monkeypatch, tmp_path):
    """正文段落→instance 候选：source_ref 含 #para、pending、携对齐结果/分数/分组键。"""
    job = _run(db, monkeypatch, tmp_path)
    assert job.status == "reviewing"

    cands = _candidates(db, job)
    instances = [c for c in cands if c.candidate_kind == "instance"]
    assert len(instances) == 2  # 两个业务实体段（PARA_ENTITY / PARA_MULTI）

    for c in instances:
        assert c.source_ref.endswith("#para")          # 溯源回链（FR-005）
        assert c.review_status == "pending"            # 不自动断言（FR-010）
        assert c.alignment_result == "new"             # align_entity 已跑（空个体→new）
        assert c.match_score is not None
        assert c.group_key                             # _compute_group_key 复用


def test_prose_coexists_with_action_candidate(db, monkeypatch, tmp_path):
    """同源 Word：prose instance 候选与 parse_action_from_text 的 Action 候选并存。"""
    job = _run(db, monkeypatch, tmp_path)
    cands = _candidates(db, job)

    actions = [c for c in cands if c.candidate_kind == "action"]
    assert len(actions) == 1
    assert actions[0].action_conditions["precondition"]
    assert "必须" in actions[0].action_conditions["obligation"]
    # 并存：Action 段不抑制 prose，prose 段不抑制 Action。
    assert any(c.candidate_kind == "instance" for c in cands)


def test_prose_backfills_iri_keys_and_aggregates_multivalue(db, monkeypatch, tmp_path):
    """label→IRI 回填；同标签多命中聚为 list（FR-006/011）。"""
    job = _run(db, monkeypatch, tmp_path)
    cands = _candidates(db, job)
    instances = [c for c in cands if c.candidate_kind == "instance"]

    # 单值段：标签经 label_to_iri 落 IRI 键。
    single = next(c for c in instances if c.extracted_properties.get(IRI_DOSAGE) == "片剂")
    assert single.extracted_properties[IRI_ACTIVE] == "化合物A"
    assert single.extracted_properties[IRI_SPEC] == "10mg"

    # 多值段：同一活性成分标签多命中 → list。
    multi = next(c for c in instances
                 if isinstance(c.extracted_properties.get(IRI_ACTIVE), list))
    assert multi.extracted_properties[IRI_ACTIVE] == ["化合物X", "化合物Y"]
