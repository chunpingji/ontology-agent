"""文档标注服务单测（document_annotator）。

覆盖三阶段标注的回填与降级：
- ``_type_and_filter_spans``：候选 span 经上下文窗口归类后 label/className 取匹配类、
  低于阈值丢弃；
- ``_span_with_context``：span 前后上下文拼接；
- ``_extract_property_triples``：对已归类 span 以类数据属性标签跑 GLiNER 抽属性值；
- ``_warnings``：仅当确有文本走 NER 时上浮 get_class_properties domain bug；
- ``annotate_excel``：GLiNER 不可用时结构完整、标注为空（优雅降级），并携带告警。

注入确定性桩嵌入器 + 以真实 openpyxl 构造样本，纯本地、零云端。
"""

from __future__ import annotations

import openpyxl

from app.services.extraction import document_annotator, ontology_typer
import pytest

from app.services.extraction.document_annotator import (
    _extract_property_triples,
    _is_short_span,
    _property_schema_for_class,
    _span_with_context,
    _type_and_filter_spans,
    _warnings,
    annotate_excel,
)
from app.services.extraction.ontology_typer import GET_CLASS_PROPERTIES_DOMAIN_BUG


class _Node:
    def __init__(self, iri, name, label, children=None):
        self.iri = iri
        self.name = name
        self.label = label
        self.children = children or []


class _Module:
    def __init__(self, key):
        self.key = key


class _FakeEngine:
    def __init__(self, roots, dp_classes, domain_props=None):
        self._roots = roots
        self._dp = dp_classes
        self._domain_props = domain_props or {}

    def get_modules(self):
        return [_Module("m")]

    def get_class_hierarchy(self, key):
        return list(self._roots)

    def data_property_domain_classes(self):
        return list(self._dp)

    def data_property_labels(self):
        return []

    def get_data_properties_by_domain(self, class_iri):
        return list(self._domain_props.get(class_iri, []))


class _FakeEmbedder:
    def __init__(self, vectors):
        self._vectors = vectors

    def is_available(self):
        return True

    def embed_many(self, texts):
        return None

    def embed(self, text):
        return self._vectors.get(text)


_PARENT = _Node("iri#drug", "Drug", "药物制剂")
_CHILD = _Node("iri#sterile", "SterilePowder", "无菌粉针剂")
_PARENT.children = [_CHILD]
_VECTORS = {
    "药物制剂": [0.99, 0.1410674],
    "无菌粉针剂": [0.97, 0.2431049],
    "无菌粉针": [1.0, 0.0],
    "随机噪声词": [0.0, 1.0],
    # 上下文窗口拼接后的 key（_span_with_context 产出）
    "无菌粉针生产线": [1.0, 0.0],
    "随机噪声词测试": [0.0, 1.0],
}


def _engine():
    return _FakeEngine([_PARENT], dp_classes=[])


def _stub_embedder(monkeypatch):
    ontology_typer._index_cache.clear()
    ontology_typer._seed_cache.clear()
    monkeypatch.setattr(ontology_typer, "get_embedder", lambda: _FakeEmbedder(_VECTORS))


# --- _span_with_context -----------------------------------------------------
def test_span_with_context_middle():
    """span 在段中间 → 前后各取窗口拼接。"""
    text = "前缀文本无菌粉针后缀文本"
    # window=2 → 取 span 前 2 字符 + span + 后 2 字符
    assert _span_with_context(text, 4, 8, window=2) == "文本无菌粉针后缀"


def test_span_with_context_at_start():
    """span 在段首 → 无前缀。"""
    assert _span_with_context("无菌粉针后缀", 0, 4, window=10) == "无菌粉针后缀"


def test_span_with_context_at_end():
    """span 在段尾 → 无后缀。"""
    assert _span_with_context("前缀无菌粉针", 2, 6, window=10) == "前缀无菌粉针"


def test_span_with_context_covers_full():
    """span 覆盖整段 → 返回原段文本。"""
    assert _span_with_context("无菌粉针", 0, 4, window=10) == "无菌粉针"


# --- _type_and_filter_spans -------------------------------------------------
def test_type_and_filter_assigns_class_and_drops(monkeypatch):
    """候选 span 归类：命中类覆盖 label/className，低于阈值的 span 被丢弃，偏移保留。"""
    _stub_embedder(monkeypatch)
    all_spans = [
        [{"start": 0, "end": 4, "text": "无菌粉针", "label": "seed", "score": 0.9}],
        [{"start": 0, "end": 5, "text": "随机噪声词", "label": "seed", "score": 0.8}],
        [],
    ]
    segment_texts = ["无菌粉针生产线", "随机噪声词测试", ""]
    out = _type_and_filter_spans(all_spans, segment_texts, _engine())
    assert len(out) == 3
    assert out[2] == []
    assert out[1] == []                            # 噪声 span 被丢弃
    assert len(out[0]) == 1
    sp = out[0][0]
    assert sp["label"] == "无菌粉针剂"            # 最具体类覆盖种子 label
    assert sp["className"] == "无菌粉针剂"        # className=label → 前端按类型着色
    assert sp["iri"] == "iri#sterile"
    assert (sp["start"], sp["end"], sp["text"]) == (0, 4, "无菌粉针")


def test_type_and_filter_degrades_without_embedder(monkeypatch):
    """嵌入器不可用 → 所有 span 丢弃（仅保留段结构）。"""
    ontology_typer._index_cache.clear()
    ontology_typer._seed_cache.clear()
    monkeypatch.setattr(ontology_typer, "get_embedder", lambda: None)
    all_spans = [[{"start": 0, "end": 4, "text": "无菌粉针", "label": "seed", "score": 0.9}]]
    assert _type_and_filter_spans(all_spans, ["无菌粉针"], _engine()) == [[]]


# --- _warnings --------------------------------------------------------------
def test_warnings_only_when_text_present():
    """有文本走 NER → 上浮 domain bug 告警；无文本 → 静默。"""
    assert _warnings(["x"]) == [GET_CLASS_PROPERTIES_DOMAIN_BUG]
    assert _warnings([]) == []
    assert _warnings(["", ""]) == []


# --- annotate_excel (优雅降级) ----------------------------------------------
def test_annotate_excel_structure_and_warning_without_gliner(tmp_path, monkeypatch):
    """GLiNER 不可用 → 行/列结构完整、标注为空，但仍记录 domain bug 告警。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)
    path = tmp_path / "a.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["设备名称", "备注"])
    ws.append(["压片机A", "额定功率15kW"])
    wb.save(path)

    data, warnings, triples, _ckpt = annotate_excel(path, _engine())

    assert data["headers"] == ["设备名称", "备注"]
    assert len(data["rows"]) == 1
    cell = data["rows"][0]["设备名称"]
    assert cell["value"] == "压片机A"
    assert cell["annotations"] == []                       # 无 GLiNER → 无标注
    assert triples == []                                   # 无 GLiNER → 无三元组
    assert warnings == [GET_CLASS_PROPERTIES_DOMAIN_BUG]


# --- _property_schema_for_class -----------------------------------------------
def test_property_schema_for_class():
    """从 engine 的 domain 反查结果派生标签集。"""
    eng = _FakeEngine([], dp_classes=[], domain_props={
        "iri#sterile": [
            {"iri": "iri#batchNo", "name": "batchNo", "label": "批号"},
            {"iri": "iri#pde", "name": "pdeValue", "label": "PDE值"},
        ],
    })
    schema = _property_schema_for_class(eng, "iri#sterile")
    assert schema["labels"] == ["批号", "PDE值"]
    assert schema["label_to_iri"] == {"批号": "iri#batchNo", "PDE值": "iri#pde"}


def test_property_schema_empty_for_unknown_class():
    """未知类 → 空 schema。"""
    schema = _property_schema_for_class(_engine(), "iri#nonexistent")
    assert schema == {"labels": [], "label_to_iri": {}}


# --- _extract_property_triples ------------------------------------------------
class _FakeGLiNerExtractor:
    """桩 GLiNER：对任何文本返回预设结果（span 格式，供 batch 调用）。"""
    def __init__(self, results=None):
        self._results = results or {}

    def is_available(self):
        return True

    def extract_text(self, text, labels, threshold=None):
        return dict(self._results)

    def extract_batch_with_spans(self, texts, labels, threshold=None):
        allowed = set(labels)
        spans = []
        for label, value in self._results.items():
            if label in allowed:
                v = str(value)
                spans.append({"start": 0, "end": len(v), "text": v, "label": label, "score": 0.9})
        return [list(spans) for _ in texts]


def test_extract_property_triples_basic(monkeypatch):
    """Stage 3：已归类 span + 类有数据属性 → 抽取属性三元组。"""
    _stub_embedder(monkeypatch)
    fake_gliner = _FakeGLiNerExtractor(results={"批号": "ABC-123"})
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: fake_gliner)

    eng = _FakeEngine([_PARENT], dp_classes=[], domain_props={
        "iri#sterile": [
            {"iri": "iri#batchNo", "name": "batchNo", "label": "批号"},
        ],
    })

    typed_spans = [
        [{"start": 0, "end": 4, "text": "无菌粉针", "label": "无菌粉针剂",
          "className": "无菌粉针剂", "score": 0.97, "iri": "iri#sterile"}],
    ]
    segment_texts = ["无菌粉针生产线批号ABC-123"]

    triples = _extract_property_triples(typed_spans, segment_texts, eng)
    assert len(triples) == 1
    t = triples[0]
    assert t["entity_text"] == "无菌粉针"
    assert t["entity_class_iri"] == "iri#sterile"
    assert t["entity_class_label"] == "无菌粉针剂"
    assert t["segment_index"] == 0
    assert len(t["properties"]) == 1
    assert t["properties"][0] == {"iri": "iri#batchNo", "label": "批号", "value": "ABC-123"}


def test_extract_property_triples_no_extractor(monkeypatch):
    """GLiNER 不可用 → 空三元组列表。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)
    triples = _extract_property_triples(
        [[{"start": 0, "end": 4, "text": "无菌粉针", "label": "X",
           "className": "X", "score": 0.9, "iri": "iri#x"}]],
        ["无菌粉针"],
        _engine(),
    )
    assert triples == []


def test_extract_property_triples_no_properties(monkeypatch):
    """类无数据属性 → 实体仍出现在三元组中但 properties 为空列表。"""
    _stub_embedder(monkeypatch)
    fake_gliner = _FakeGLiNerExtractor()
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: fake_gliner)

    typed_spans = [
        [{"start": 0, "end": 4, "text": "无菌粉针", "label": "无菌粉针剂",
          "className": "无菌粉针剂", "score": 0.97, "iri": "iri#sterile"}],
    ]
    triples = _extract_property_triples(typed_spans, ["无菌粉针"], _engine())
    assert len(triples) == 1
    assert triples[0]["properties"] == []


def test_extract_property_triples_multiple_classes(monkeypatch):
    """多个不同类的实体 → 各自用自己的属性标签抽取。"""
    _stub_embedder(monkeypatch)
    extract_calls = []

    class _TrackingExtractor:
        def is_available(self):
            return True

        def extract_text(self, text, labels, threshold=None):
            extract_calls.append(labels)
            return {}

        def extract_batch_with_spans(self, texts, labels, threshold=None):
            extract_calls.append(labels)
            return [[] for _ in texts]

    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: _TrackingExtractor())

    eng = _FakeEngine([_PARENT], dp_classes=[], domain_props={
        "iri#drug": [{"iri": "iri#name", "name": "name", "label": "通用名"}],
        "iri#sterile": [{"iri": "iri#batch", "name": "batch", "label": "批号"}],
    })

    typed_spans = [
        [{"start": 0, "end": 2, "text": "药物", "label": "药物制剂",
          "className": "药物制剂", "score": 0.9, "iri": "iri#drug"}],
        [{"start": 0, "end": 4, "text": "无菌粉针", "label": "无菌粉针剂",
          "className": "无菌粉针剂", "score": 0.97, "iri": "iri#sterile"}],
    ]

    triples = _extract_property_triples(typed_spans, ["药物制剂A", "无菌粉针生产线"], eng)
    assert len(triples) == 2
    assert set(t["entity_class_iri"] for t in triples) == {"iri#drug", "iri#sterile"}
    assert ["通用名"] in extract_calls
    assert ["批号"] in extract_calls


# --- _is_short_span -----------------------------------------------------------

@pytest.mark.parametrize("text,expected", [
    ("设备", True),       # 2 CJK ≤ 3
    ("离心机", True),      # 3 CJK ≤ 3
    ("API", True),        # 3 ASCII ≤ 4
    ("HPLC", True),       # 4 ASCII ≤ 4
    ("反应釜设备", False),  # 5 CJK > 3
    ("HRS-1234", False),  # 8 chars > 4
    ("无菌粉针", False),   # 4 CJK > 3
    ("pH", True),         # 2 ASCII ≤ 4
])
def test_is_short_span(text, expected):
    assert _is_short_span(text) is expected


# --- 双 pass 归类策略 ----------------------------------------------------------

def test_dual_pass_short_span_raw_match(monkeypatch):
    """短 span "设备" 直接走 Pass-1 raw text 匹配，不被上下文拉偏到 "清洗设备"。"""
    equip = _Node("iri#equip", "Equipment", "设备")
    cleaning_equip = _Node("iri#clean_equip", "CleaningEquipment", "清洗设备")
    equip.children = [cleaning_equip]
    eng = _FakeEngine([equip], dp_classes=[])

    vectors = {
        "设备": [1.0, 0.0],
        "清洗设备": [0.6, 0.8],
        # 上下文文本 "清洗设备残留" 会偏向清洗设备
        "清洗设备残留": [0.5, 0.87],
    }
    ontology_typer._index_cache.clear()
    ontology_typer._seed_cache.clear()
    monkeypatch.setattr(ontology_typer, "get_embedder", lambda: _FakeEmbedder(vectors))

    all_spans = [[{"start": 2, "end": 4, "text": "设备", "label": "seed", "score": 0.9}]]
    segment_texts = ["清洗设备残留"]
    out = _type_and_filter_spans(all_spans, segment_texts, eng)
    assert len(out[0]) == 1
    assert out[0][0]["label"] == "设备"  # raw match, not "清洗设备"


def test_dual_pass_opaque_span_uses_context(monkeypatch):
    """不透明 span "HRS-1234" Pass-1 无高置信命中 → Pass-2 用上下文匹配到正确类。"""
    drug = _Node("iri#drug", "Drug", "药物制剂")
    equip = _Node("iri#equip", "Equipment", "设备")
    eng = _FakeEngine([drug, equip], dp_classes=[])

    vectors = {
        "药物制剂": [1.0, 0.0],
        "设备": [0.0, 1.0],
        # raw "HRS-1234" 无方向 → 两类余弦都不够 0.85
        "HRS-1234": [0.4, 0.3],
        # 上下文 "原料药HRS-1234的合成路线描述" 拉向药物制剂（_CONTEXT_WINDOW=15 覆盖全段）
        "原料药HRS-1234的合成路线描述": [0.85, 0.1],
    }
    ontology_typer._index_cache.clear()
    ontology_typer._seed_cache.clear()
    monkeypatch.setattr(ontology_typer, "get_embedder", lambda: _FakeEmbedder(vectors))

    all_spans = [[{"start": 3, "end": 11, "text": "HRS-1234", "label": "seed", "score": 0.9}]]
    segment_texts = ["原料药HRS-1234的合成路线描述"]
    out = _type_and_filter_spans(all_spans, segment_texts, eng)
    assert len(out[0]) == 1
    assert out[0][0]["label"] == "药物制剂"  # context-assisted match


# --- 013: parse_word_to_tiptap — 结构化解析（跳过 NER，engine=None 安全）---------


def _node_types(node, acc):
    if isinstance(node, dict):
        if node.get("type"):
            acc.add(node["type"])
        for c in node.get("content") or []:
            _node_types(c, acc)
    return acc


def _mark_types(node, acc):
    if isinstance(node, dict):
        for m in node.get("marks") or []:
            if isinstance(m, dict) and m.get("type"):
                acc.add(m["type"])
        for c in node.get("content") or []:
            _mark_types(c, acc)
    return acc


def test_parse_word_to_tiptap_structure_without_ner(tmp_path):
    """结构化解析产出忠于原文的 tiptap（标题 + 表格），零 entity-annotation 标注、
    engine=None 安全（不触碰 GLiNER/嵌入权重）——即 013 后台结构化解析的契约。"""
    import docx
    from app.services.extraction.document_annotator import parse_word_to_tiptap

    path = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_heading("评估对象", level=1)
    doc.add_paragraph("本品为XX注射液，剂型为注射剂。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "设备编号"
    table.rows[0].cells[1].text = "设备名称"
    table.rows[1].cells[0].text = "CT64201"
    table.rows[1].cells[1].text = "压片机A"
    doc.save(str(path))

    result = parse_word_to_tiptap(path)

    # 忠于结构：doc 根 + 标题节点 + 表格节点都在。
    assert result["type"] == "doc"
    types = _node_types(result, set())
    assert "heading" in types
    assert "table" in types
    # 结构化解析跳过 NER：绝无 entity-annotation mark（格式 mark 可合法存在）。
    assert "entity-annotation" not in _mark_types(result, set())
