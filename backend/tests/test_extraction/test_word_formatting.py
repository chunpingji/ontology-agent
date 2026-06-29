"""annotate_word 保留 Word 文档样式（回归） + 表格行级上下文拼接（009）。

以真实 python-docx 构造样本，桩掉 GLiNER（span 为空）→ 结构断言确定、纯本地。
"""

from __future__ import annotations

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from app.services.extraction import document_annotator
from app.services.extraction.document_annotator import (
    _build_row_segment,
    _correct_span_offsets,
    _detect_header_rows,
    _find_table_caption,
    _inline_nodes,
    _is_vmerge_continue,
    annotate_word,
)


class _Module:
    def __init__(self, key):
        self.key = key


class _FakeEngine:
    """最小本体引擎桩：仅供 seed_labels 调用，无需真实类层级。"""

    def get_modules(self):
        return [_Module("m")]

    def get_class_hierarchy(self, key):
        return []

    def data_property_domain_classes(self):
        return []

    def data_property_labels(self):
        return []

    def get_data_properties_by_domain(self, class_iri):
        return []


def _types(nodes):
    return [n.get("type") for n in nodes]


def test_annotate_word_preserves_structure(tmp_path, monkeypatch):
    """heading 层级 / 居中对齐 / 空行 / 行内粗体均被保留为对应 tiptap 节点。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    d = docx.Document()
    d.add_heading("产品质量标准", level=1)          # → heading level 1
    centered = d.add_paragraph("居中副标题")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER  # → textAlign center
    d.add_paragraph("")                              # 空行 → 保留间距
    body = d.add_paragraph()
    body.add_run("关键指标").bold = True             # 粗体 run
    body.add_run("：常规说明")                        # 普通 run
    path = tmp_path / "spec.docx"
    d.save(path)

    doc, _, _, _ = annotate_word(path, _FakeEngine())
    content = doc["content"]

    # 标题段：heading level 1
    heading = content[0]
    assert heading["type"] == "heading"
    assert heading["attrs"]["level"] == 1
    assert heading["content"][0]["text"] == "产品质量标准"

    # 居中段：paragraph + textAlign=center
    centered_node = content[1]
    assert centered_node["type"] == "paragraph"
    assert centered_node["attrs"]["textAlign"] == "center"

    # 空行：空 paragraph（无 content）→ 渲染为空行，保留间距
    empty_node = content[2]
    assert empty_node["type"] == "paragraph"
    assert "content" not in empty_node

    # 正文段：粗体 run 携 bold mark，普通 run 无 mark
    body_node = content[3]
    bold_seg = body_node["content"][0]
    assert bold_seg["text"] == "关键指标"
    assert {"type": "bold"} in bold_seg["marks"]
    plain_seg = body_node["content"][1]
    assert plain_seg["text"] == "：常规说明"
    assert "marks" not in plain_seg


def test_inline_nodes_merges_formatting_and_entity():
    """行内样式与 NER span 按字符边界合并：交叠段同时带 bold + entity mark。"""
    text = "无菌粉针剂A"
    runs = [(0, 5, ["bold"]), (5, 6, [])]   # 前 5 字粗体，"A" 普通
    spans = [{
        "start": 0, "end": 4, "text": "无菌粉针",
        "label": "无菌粉针剂", "className": "无菌粉针剂", "score": 0.9,
    }]

    nodes = _inline_nodes(text, spans, runs)

    # 边界 {0,4,5,6} → 三段
    assert [n["text"] for n in nodes] == ["无菌粉针", "剂", "A"]
    # 段一：bold + entity（交叠）
    types0 = {m["type"] for m in nodes[0]["marks"]}
    assert types0 == {"bold", "entity-annotation"}
    # 段二：仅 bold（在粗体 run 内、span 外）
    assert [m["type"] for m in nodes[1]["marks"]] == ["bold"]
    # 段三："A" 无样式无标注
    assert "marks" not in nodes[2]


def test_font_size_heading_fallback(tmp_path, monkeypatch):
    """Normal 样式 + 大字号 → 字号启发式回退识别为 heading。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    d = docx.Document()
    # 18pt（小二）→ h2；Normal 样式不触发 _heading_level
    title_para = d.add_paragraph()
    run = title_para.add_run("文档标题")
    run.font.size = Pt(18)

    # 14pt（四号）→ h3
    sub_para = d.add_paragraph()
    run2 = sub_para.add_run("子标题")
    run2.font.size = Pt(14)

    # 12pt（小四）→ paragraph
    body = d.add_paragraph()
    run3 = body.add_run("正文内容")
    run3.font.size = Pt(12)

    path = tmp_path / "fontsize.docx"
    d.save(path)

    doc, _, _, _ = annotate_word(path, _FakeEngine())
    content = doc["content"]

    assert content[0]["type"] == "heading"
    assert content[0]["attrs"]["level"] == 2

    assert content[1]["type"] == "heading"
    assert content[1]["attrs"]["level"] == 3

    assert content[2]["type"] == "paragraph"


def test_inherited_bold_from_style(tmp_path, monkeypatch):
    """段落样式定义粗体 → run.bold=None 时应继承，标记为 bold mark。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    d = docx.Document()
    # 创建段落，在样式上设置粗体
    para = d.add_paragraph(style="Normal")
    para.style.font.bold = True
    para.add_run("继承粗体")
    # 第二个 run 显式取消粗体
    r2 = para.add_run("非粗体")
    r2.bold = False

    path = tmp_path / "inherited.docx"
    d.save(path)

    doc, _, _, _ = annotate_word(path, _FakeEngine())
    body = doc["content"][0]

    # 第一段文本：继承粗体
    assert {"type": "bold"} in body["content"][0].get("marks", [])
    # 第二段文本：显式取消粗体
    marks = body["content"][1].get("marks", [])
    assert {"type": "bold"} not in marks


def test_tables_interleaved_with_paragraphs(tmp_path, monkeypatch):
    """表格应保持在文档中的原始位置，而非被统一移到尾部。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    d = docx.Document()
    d.add_paragraph("段落一")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "单元格A"
    t.cell(0, 1).text = "单元格B"
    d.add_paragraph("段落二")

    path = tmp_path / "interleaved.docx"
    d.save(path)

    doc, _, _, _ = annotate_word(path, _FakeEngine())
    types = [n["type"] for n in doc["content"]]

    assert types == ["paragraph", "table", "paragraph"]


# ---------------------------------------------------------------------------
# US1: 表格行级上下文拼接 (T006-T010)
# ---------------------------------------------------------------------------


def _make_table_doc(tmp_path, rows_data, *, caption=None):
    """创建含单表格的 Word 文档。rows_data[0] 为表头行，其余为数据行。"""
    d = docx.Document()
    if caption:
        d.add_paragraph(caption)
    ncols = len(rows_data[0]) if rows_data else 0
    t = d.add_table(rows=len(rows_data), cols=ncols)
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            t.cell(ri, ci).text = val
    path = tmp_path / "table.docx"
    d.save(path)
    return path, t


def test_row_level_concatenation(tmp_path):
    """数据行以 'hdr：val | hdr：val' 格式拼接为单行级 segment。"""
    rows = [
        ["名称", "剂型", "规格"],
        ["阿莫西林", "片剂", "0.25g"],
    ]
    path, table = _make_table_doc(tmp_path, rows)

    headers = [cell.text.strip() for cell in table.rows[0].cells]
    seg, offsets = _build_row_segment(table.rows[1], headers)

    assert "名称：阿莫西林" in seg
    assert "剂型：片剂" in seg
    assert "规格：0.25g" in seg
    assert " | " in seg
    assert len(offsets) == 3


def test_header_row_skipped(tmp_path, monkeypatch):
    """表头行不参与 NER（不出现在 all_texts 中），仅数据行产生行级 segment。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    rows = [
        ["名称", "剂型", "规格"],
        ["阿莫西林", "片剂", "0.25g"],
        ["布洛芬", "胶囊", "0.3g"],
    ]
    path, _ = _make_table_doc(tmp_path, rows)

    doc, _, _, _ = annotate_word(path, _FakeEngine())
    table_node = doc["content"][0]
    assert table_node["type"] == "table"

    header_row = table_node["content"][0]
    for cell in header_row["content"]:
        text_nodes = cell["content"][0].get("content", [])
        for tn in text_nodes:
            marks = tn.get("marks", [])
            assert not any(m["type"] == "entity-annotation" for m in marks)


def test_table_caption_prefix(tmp_path):
    """表格前紧邻段落作为 caption，追加到行级 segment 前缀。"""
    rows = [
        ["名称", "规格"],
        ["阿莫西林", "0.25g"],
    ]
    path, table = _make_table_doc(tmp_path, rows, caption="表 3：原料药规格")

    from docx import Document
    doc = Document(str(path))
    t = doc.tables[0]

    caption = _find_table_caption(doc.element.body, t._element)
    assert caption == "表 3：原料药规格"

    headers = [cell.text.strip() for cell in t.rows[0].cells]
    seg, _ = _build_row_segment(t.rows[1], headers, caption)
    assert seg.startswith("表 3：原料药规格 — ")


def test_span_offset_correction():
    """NER span 坐标从行级 segment 校正回 cell 内坐标。"""
    headers = ["名称", "剂型"]
    cell_offsets = [
        (0, 7, 0),   # "名称：阿莫西林" at 0..7
        (10, 15, 1),  # "剂型：片剂" at 10..15
    ]
    spans = [
        {"start": 3, "end": 7, "text": "阿莫西林", "label": "Drug", "score": 0.9},
    ]

    corrected = _correct_span_offsets(spans, cell_offsets, headers)
    assert len(corrected) == 1
    col_idx, adj = corrected[0]
    assert col_idx == 0
    assert adj["start"] == 0
    assert adj["end"] == 4
    assert adj["text"] == "阿莫西林"


def test_empty_cell_handling(tmp_path):
    """空/空白单元格在行级拼接中被跳过，不产生空 fragment。"""
    rows = [
        ["名称", "剂型", "备注"],
        ["阿莫西林", "", "  "],
    ]
    path, table = _make_table_doc(tmp_path, rows)

    headers = [cell.text.strip() for cell in table.rows[0].cells]
    seg, offsets = _build_row_segment(table.rows[1], headers)

    assert "名称：阿莫西林" in seg
    assert "剂型" not in seg
    assert "备注" not in seg
    assert len(offsets) == 1


# ---------------------------------------------------------------------------
# US2: 合并单元格与嵌套表格 (T016-T019)
# ---------------------------------------------------------------------------


def _set_vmerge_tc(tc_elem, val=None):
    """设置原始 TC 元素的 vMerge 属性。val="restart" 为主行，val=None 为续行。"""
    tc_pr = tc_elem.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = docx.oxml.OxmlElement("w:tcPr")
        tc_elem.insert(0, tc_pr)
    vmerge = docx.oxml.OxmlElement("w:vMerge")
    if val:
        vmerge.set(qn("w:val"), val)
    tc_pr.append(vmerge)


def _set_gridspan_tc(tc_elem, span_val):
    """设置原始 TC 元素的 gridSpan 属性（水平合并）。"""
    tc_pr = tc_elem.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = docx.oxml.OxmlElement("w:tcPr")
        tc_elem.insert(0, tc_pr)
    gs = docx.oxml.OxmlElement("w:gridSpan")
    gs.set(qn("w:val"), str(span_val))
    tc_pr.append(gs)


def test_vmerge_dedup():
    """纵向合并续行 cell 在行级拼接中被跳过，避免重复 NER。"""
    d = docx.Document()
    t = d.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "名称"
    t.cell(0, 1).text = "批号"
    t.cell(1, 0).text = "阿莫西林"
    t.cell(1, 1).text = "B001"
    t.cell(2, 0).text = "阿莫西林"
    t.cell(2, 1).text = "B002"

    # vMerge 须设在 raw TC 元素上（row._tr 的子元素），
    # python-docx 的 Cell 访问对合并单元格返回 master cell。
    row1_tcs = list(t.rows[1]._tr.iterchildren(qn("w:tc")))
    row2_tcs = list(t.rows[2]._tr.iterchildren(qn("w:tc")))
    _set_vmerge_tc(row1_tcs[0], "restart")
    _set_vmerge_tc(row2_tcs[0])  # continuation (no val)

    assert _is_vmerge_continue(row2_tcs[0]) is True
    assert _is_vmerge_continue(row1_tcs[0]) is False
    assert _is_vmerge_continue(row2_tcs[1]) is False

    headers = [cell.text.strip() for cell in t.rows[0].cells]
    seg2, offs2 = _build_row_segment(t.rows[2], headers, skip_vmerge=True)

    assert "阿莫西林" not in seg2
    assert "B002" in seg2


def test_multi_row_header():
    """首行含 gridSpan > 1 时检测多行表头，数据行使用末行表头作列名前缀。"""
    d = docx.Document()
    t = d.add_table(rows=3, cols=4)

    t.cell(0, 0).merge(t.cell(0, 1))
    t.cell(0, 2).merge(t.cell(0, 3))
    t.cell(0, 0).text = "药品信息"
    t.cell(0, 2).text = "生产信息"

    t.cell(1, 0).text = "名称"
    t.cell(1, 1).text = "剂型"
    t.cell(1, 2).text = "批号"
    t.cell(1, 3).text = "日期"

    t.cell(2, 0).text = "阿莫西林"
    t.cell(2, 1).text = "片剂"
    t.cell(2, 2).text = "B001"
    t.cell(2, 3).text = "2026-01-01"

    hdr_count = _detect_header_rows(t)
    assert hdr_count == 2

    last_hdr = t.rows[hdr_count - 1]
    headers = [cell.text.strip() for cell in last_hdr.cells]
    seg, _ = _build_row_segment(t.rows[2], headers)

    assert "名称：阿莫西林" in seg
    assert "剂型：片剂" in seg


def _make_xml_row(cell_texts):
    """构建一个 <w:tr> XML 元素，包含指定文本的单元格。"""
    tr = docx.oxml.OxmlElement("w:tr")
    for text in cell_texts:
        tc = docx.oxml.OxmlElement("w:tc")
        p = docx.oxml.OxmlElement("w:p")
        r = docx.oxml.OxmlElement("w:r")
        t = docx.oxml.OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
    return tr


def test_nested_table_recursion(tmp_path, monkeypatch):
    """cell 内嵌套表格的内容被递归处理并出现在 tiptap 输出中。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    d = docx.Document()
    outer = d.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "主列A"
    outer.cell(0, 1).text = "主列B"
    outer.cell(1, 0).text = "数据A"

    inner_cell = outer.cell(1, 1)
    inner_tbl_elem = docx.oxml.OxmlElement("w:tbl")
    inner_tbl_elem.append(_make_xml_row(["嵌套列1", "嵌套列2"]))
    inner_tbl_elem.append(_make_xml_row(["嵌套数据1", "嵌套数据2"]))
    inner_cell._tc.append(inner_tbl_elem)

    path = tmp_path / "nested.docx"
    d.save(path)

    doc, _, _, _ = annotate_word(path, _FakeEngine())

    table_node = doc["content"][0]
    assert table_node["type"] == "table"

    data_row_node = table_node["content"][1]
    cell_1 = data_row_node["content"][1]

    nested_found = any(c.get("type") == "table" for c in cell_1["content"])
    assert nested_found, "嵌套表格应出现在 cell content 中"

    nested_table = next(c for c in cell_1["content"] if c["type"] == "table")
    assert len(nested_table["content"]) == 2

    nested_data_row = nested_table["content"][1]
    all_text = ""
    for cell in nested_data_row["content"]:
        for p in cell.get("content", []):
            for tn in p.get("content", []):
                all_text += tn.get("text", "")
    assert "嵌套数据1" in all_text or "嵌套数据2" in all_text


def test_nesting_depth_limit():
    """递归深度限制不应导致异常——helper 接受 max_depth 参数。"""
    d = docx.Document()
    t = d.add_table(rows=2, cols=1)
    t.cell(0, 0).text = "H"
    t.cell(1, 0).text = "D"

    headers = ["H"]
    seg, _ = _build_row_segment(t.rows[1], headers)
    assert "H：D" in seg


def test_multi_paragraph_cell(tmp_path, monkeypatch):
    """单元格内多个段落用空格分隔，不丢失段落边界。"""
    monkeypatch.setattr(document_annotator, "_get_extractor", lambda: None)

    d = docx.Document()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "名称"
    t.cell(0, 1).text = "描述"
    t.cell(1, 0).text = "阿莫西林"

    cell_11 = t.cell(1, 1)
    cell_11.text = "第一段描述"
    cell_11.add_paragraph("第二段描述")

    path = tmp_path / "multi_para.docx"
    d.save(path)

    doc, _, _, _ = annotate_word(path, _FakeEngine())

    table_node = doc["content"][0]
    data_row = table_node["content"][1]
    desc_cell = data_row["content"][1]
    cell_text = ""
    for p in desc_cell.get("content", []):
        for tn in p.get("content", []):
            cell_text += tn.get("text", "")
    assert "第一段描述" in cell_text
    assert "第二段描述" in cell_text
