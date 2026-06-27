"""集成测试：本地 NER 不可用时优雅降级（008 Polish T026）。

覆盖 [contracts/offline-extraction-invariants.md](../../../specs/008-gliner-ner-extraction/contracts/offline-extraction-invariants.md)
O6–O8 / FR-012 / SC-006——这里走**真实** `settings.gliner_extraction_enabled=False`
门控（`get_gliner_extractor()` 返回 `None`），与 US2/US3 直接注入桩的测试互补：

- O6 作业不失败：Word/Excel 输入在 NER 关闭时仍正常完成（`status=reviewing`）。
- O7 结构化零回归：prose 候选为空、Excel 不富化、候选数 = 行数、暂存仍清除。
- O8 运维可见、用户无感：不可用记 `logger.warning`，作业无 `error_message`、不标 `degraded`。

纯本地、零云端（`llm_cloud_enabled=False`、无 Key → 离线正常态、非降级）。
"""

from __future__ import annotations

import asyncio
import logging
from types import SimpleNamespace

import docx
import openpyxl

from app.config import settings
from app.models.extraction import ExtractionCandidate, ExtractionConfig, ExtractionJob
from app.services.extraction.gliner_extractor import GlinerExtractor, get_gliner_extractor
from app.services.extraction.pipeline import run_extraction_pipeline

DRUG = "https://ontology.pharma-gmp.cn/slpra/drug/DrugProduct"
EQUIP = "http://slpra.org/equipment#Equipment"
IRI_ACTIVE = f"{DRUG}#activeIngredient"
IRI_DOSAGE = f"{DRUG}#dosageForm"
IRI_ID = f"{EQUIP}#equipmentID"
IRI_NAME = f"{EQUIP}#equipmentName"
IRI_POWER = f"{EQUIP}#ratedPower"

PARA_ENTITY = "本品活性成分为化合物A，剂型为片剂。"
PARA_ACTION = "若产品为高致敏药品，则必须执行专用化生产。"
NOTE = "额定功率15kW，名称压片机甲。"


class _Engine:
    """假本体引擎：派生非空标签集（确保 prose 为空是因 NER 关、而非 schema 空）。"""

    def __init__(self, data_properties):
        self._dp = data_properties

    def get_class_detail(self, iri):
        return SimpleNamespace(data_properties=self._dp)

    def get_individuals(self, iri):
        return []

    def __getattr__(self, name):
        def _noop(*a, **k):
            return None
        return _noop


def _drug_engine():
    return _Engine([
        {"iri": IRI_ACTIVE, "name": "activeIngredient", "label": "活性成分", "range": ["string"]},
        {"iri": IRI_DOSAGE, "name": "dosageForm", "label": "剂型", "range": ["string"]},
    ])


def _equip_engine():
    return _Engine([
        {"iri": IRI_NAME, "name": "equipmentName", "label": "设备名称", "range": ["string"]},
        {"iri": IRI_POWER, "name": "ratedPower", "label": "额定功率", "range": ["string"]},
    ])


def _offline(monkeypatch):
    """air-gap 默认：云端关、无 Key、本地 NER 关（真实门控 → get_gliner_extractor=None）。"""
    monkeypatch.setattr(settings, "llm_cloud_enabled", False)
    monkeypatch.setattr(settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings, "gliner_extraction_enabled", False)
    get_gliner_extractor.cache_clear()  # 清单例缓存，使新 flag 生效


def _job(db, source_type, filename):
    job = ExtractionJob(source_type=source_type, source_filename=filename,
                        source_config={}, status="pending")
    db.add(job)
    db.commit()
    db.refresh(job)
    return job


def _candidates(db, job):
    return db.query(ExtractionCandidate).filter(
        ExtractionCandidate.job_id == job.id).all()


# --- O6/O7/O8 Word：NER 关 → 仍成功、prose 空、不 degraded ---------------------
def test_word_ner_disabled_succeeds_no_prose_not_degraded(db, monkeypatch, tmp_path):
    _offline(monkeypatch)
    doc = docx.Document()
    doc.add_paragraph(PARA_ENTITY)
    doc.add_paragraph(PARA_ACTION)
    path = tmp_path / "SOP.docx"
    doc.save(path)

    cfg = ExtractionConfig(name="药物", target_class_iri=DRUG, source_type="word",
                           column_mapping={"活性成分": IRI_ACTIVE, "剂型": IRI_DOSAGE})
    db.add(cfg)
    db.commit()
    db.refresh(cfg)
    job = _job(db, "word", "SOP.docx")
    asyncio.run(run_extraction_pipeline(job, cfg, path, _drug_engine(), db))
    db.refresh(job)

    assert job.status == "reviewing"                       # O6 作业不失败
    assert job.error_message is None                       # O8 无报错
    cands = _candidates(db, job)
    # O7 prose 候选为空：无 #para 来源的 instance 候选（NER 关）。
    assert not [c for c in cands if c.candidate_kind == "instance"]
    # Action 通道与 NER 正交，照常产出（FR-005 不受降级影响）。
    assert [c for c in cands if c.candidate_kind == "action"]
    assert all(c.degraded_reason is None for c in cands)   # O8 不标 degraded


# --- O6/O7 Excel：NER 关 → 不富化、候选=行数、暂存清除、不 degraded -------------
def test_excel_ner_disabled_no_enrichment_zero_regression(db, monkeypatch, tmp_path):
    _offline(monkeypatch)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["设备编号", "设备名称", "备注"])
    ws.append(["CT64201", "压片机A", NOTE])
    ws.append(["DE64203", "包衣机B", NOTE])
    path = tmp_path / "equip.xlsx"
    wb.save(path)

    cfg = ExtractionConfig(name="设备", target_class_iri=EQUIP, source_type="excel",
                           column_mapping={"设备编号": IRI_ID, "设备名称": IRI_NAME},
                           ner_columns=["备注"])
    db.add(cfg)
    db.commit()
    db.refresh(cfg)
    job = _job(db, "excel", "equip.xlsx")
    asyncio.run(run_extraction_pipeline(job, cfg, path, _equip_engine(), db))
    db.refresh(job)

    assert job.status == "reviewing"                       # O6
    cands = _candidates(db, job)
    assert len(cands) == 2                                 # 候选数 = 行数（不另生候选）
    for c in cands:
        assert "__freetext__" not in c.extracted_properties   # O7 暂存仍清除
        assert IRI_POWER not in c.extracted_properties         # O7 未富化（NER 关）
        assert c.degraded_reason is None                       # O8 不 degraded
    # 结构化属性逐字保留。
    row1 = next(c for c in cands if c.extracted_properties.get(IRI_ID) == "CT64201")
    assert row1.extracted_properties[IRI_NAME] == "压片机A"


# --- O8 运维可见：缺包/缺权重 → is_available() False 且记 WARNING -----------------
def test_extractor_unavailable_logs_warning(caplog, monkeypatch, tmp_path):
    """真实 GlinerExtractor：缺 gliner 包或缺本地权重 → 静默降级 + WARNING（不抛出）。

    把权重路径指向不存在的本地目录（local_files_only=True）强制加载失败，从而**确定性**
    复现「不可用」契约——不依赖运行环境是否恰好缺权重（开发机已下权重时 is_available 本会
    为真）。air-gap 缺权重与此路径等价。
    """
    monkeypatch.setattr(settings, "gliner_model_path", str(tmp_path / "no_such_weights"))
    extractor = GlinerExtractor()
    with caplog.at_level(logging.WARNING, logger="app.services.extraction.gliner_extractor"):
        available = extractor.is_available()
    # 测试环境无本地权重/或无 gliner 包 → 不可用、记 WARNING、extract_text 返 {} 不抛。
    assert available is False
    assert any(r.levelno == logging.WARNING for r in caplog.records)
    assert extractor.extract_text("额定功率15kW", ["额定功率"]) == {}
