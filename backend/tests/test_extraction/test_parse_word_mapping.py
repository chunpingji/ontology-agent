"""契约测试：Word 表头确定性映射（parse_word column_mapping）。

覆盖 [contracts/parser-and-enrichment.md](../../../specs/008-gliner-ner-extraction/contracts/parser-and-enrichment.md)
P1–P4：表头命中映射→IRI 键、未命中原样保留（确定性，替代云端）；零 LLM；
`column_mapping=None` 向后兼容；段落形态不变。以真实 python-docx 构造样本，
不触任何云端——风格对齐既有 test_extraction_pipeline.py 的 docx 构造做法。
"""

from __future__ import annotations

import sys
import types

import docx

from app.services.extraction.parser import parse_word

MAPPING = {"设备编号": "equipmentID", "设备名称": "equipmentName"}


def _make_docx(path):
    """3 列表（2 已映射 + 1 未映射）+ 1 业务段落。"""
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "设备编号", "设备名称", "未映射列"
    body = table.rows[1].cells
    body[0].text, body[1].text, body[2].text = "CT64201", "压片机A", "备注内容"
    doc.add_paragraph("若设备用于高致敏药品生产，则必须执行专用化管理。")
    doc.save(path)
    return path


def _table_rows(sections):
    return [s for s in sections if s.get("type") == "table_row"]


def _paragraphs(sections):
    return [s for s in sections if s.get("type") == "paragraph"]


def test_mapped_headers_become_iri_keys(tmp_path):
    """P1：表头命中 column_mapping → 行以 IRI 为键；未命中 → 原表头键保留。"""
    path = _make_docx(tmp_path / "sop.docx")
    rows = _table_rows(parse_word(path, MAPPING))
    assert len(rows) == 1
    content = rows[0]["content"]
    assert content == {
        "equipmentID": "CT64201",
        "equipmentName": "压片机A",
        "未映射列": "备注内容",  # 未命中映射 → 原表头键容忍保留
    }


def test_no_cloud_during_word_parse(tmp_path, monkeypatch):
    """P2：Word 表头→IRI 为纯本地确定性映射，绝不触达 anthropic。"""

    class _Boom(types.ModuleType):
        def __getattr__(self, name):
            raise AssertionError(f"Word 解析不得触达 anthropic.{name}")

    monkeypatch.setitem(sys.modules, "anthropic", _Boom("anthropic"))
    path = _make_docx(tmp_path / "sop.docx")
    rows = _table_rows(parse_word(path, MAPPING))
    assert rows[0]["content"]["equipmentID"] == "CT64201"


def test_none_mapping_backward_compatible(tmp_path):
    """P3：column_mapping=None → 与改造前一致（原表头为键）。"""
    path = _make_docx(tmp_path / "sop.docx")
    rows = _table_rows(parse_word(path))
    assert rows[0]["content"] == {
        "设备编号": "CT64201",
        "设备名称": "压片机A",
        "未映射列": "备注内容",
    }


def test_paragraph_shape_unchanged(tmp_path):
    """P4：段落解析形态不变（仍产 {"type":"paragraph","content",...}）。"""
    path = _make_docx(tmp_path / "sop.docx")
    paras = _paragraphs(parse_word(path, MAPPING))
    assert len(paras) == 1
    assert paras[0]["type"] == "paragraph"
    assert paras[0]["content"] == "若设备用于高致敏药品生产，则必须执行专用化管理。"
    assert "style" in paras[0]
