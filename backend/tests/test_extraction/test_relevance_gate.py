"""相关性门控：自动/未映射抽取下，行仅在提及目标类时才落候选（防「行×类」放大）。

背景：自动抽取（POST /jobs/auto）对全部目标类逐类各跑一遍流水线；离线结构化透传
把每张表的每一行落到当前目标类，致候选被「行×类」放大约 200 倍（实测单文档 14,886）。
门控以「类标签/名称是否出现在行的键/值文本」作判定，把笛卡尔积收敛为真实相关对，
且仅在未显式配置 column_mapping 时启用——已配置即分析师声明此源映射此类，零回归。
"""

from __future__ import annotations

import asyncio
import types

import docx
import pytest

from app.config import settings
from app.models.extraction import ExtractionCandidate, ExtractionConfig, ExtractionJob
from app.services.extraction.pipeline import (
    _class_label_tokens,
    _row_mentions_class,
    run_extraction_pipeline,
)
from tests.conftest import FakeOntologyEngine

EQUIP = "http://slpra.org/equipment#Equipment"


# --- 纯函数单测 -------------------------------------------------------------

class _StubEngine(FakeOntologyEngine):
    """get_class_detail 返回带标签的 detail；其余调用沿用 FakeOntologyEngine 的 noop。"""

    def __init__(self, detail) -> None:
        super().__init__()
        self._detail = detail

    def get_class_detail(self, iri):  # noqa: ARG002
        return self._detail


def _detail(**kw):
    base = {"label_zh": None, "label_en": None, "name": None, "data_properties": []}
    base.update(kw)
    return types.SimpleNamespace(**base)


def test_class_label_tokens_collects_and_filters():
    eng = _StubEngine(_detail(label_zh="设备", label_en="Equipment", name="Equipment"))
    assert _class_label_tokens(eng, EQUIP) == {"设备", "Equipment"}

    # 单字/空白标记剔除（避免噪声过匹配）；name 为 None 不计。
    eng2 = _StubEngine(_detail(label_zh="A", label_en="  ", name="反应器"))
    assert _class_label_tokens(eng2, EQUIP) == {"反应器"}

    # 类不存在 → 空集（调用方据此放行，无从判定不误杀）。
    assert _class_label_tokens(_StubEngine(None), EQUIP) == set()


def test_row_mentions_class_keep_drop_and_passthrough():
    tokens = {"设备"}
    assert _row_mentions_class({"设备规格": "500L反应釜"}, tokens) is True   # 键命中
    assert _row_mentions_class({"x": "本设备用于压片"}, tokens) is True       # 值命中
    assert _row_mentions_class({"活性成分": "HRS-1234"}, tokens) is False     # 无命中
    assert _row_mentions_class({"活性成分": "HRS-1234"}, set()) is True       # 空标记放行


# --- 集成：端到端经 run_extraction_pipeline -------------------------------

def _mixed_docx(path):
    """两张表：一张提及「设备」、一张为药品列（不提及设备）。"""
    doc = docx.Document()
    t1 = doc.add_table(rows=2, cols=1)
    t1.rows[0].cells[0].text = "设备名称"
    t1.rows[1].cells[0].text = "压片机"
    t2 = doc.add_table(rows=2, cols=1)
    t2.rows[0].cells[0].text = "活性成分"
    t2.rows[1].cells[0].text = "HRS-1234"
    doc.save(path)
    return path


def _job(db):
    job = ExtractionJob(source_type="word", source_filename="mix.docx",
                        source_config={}, status="pending")
    db.add(job)
    db.commit()
    db.refresh(job)
    return job


def _instances(db, job):
    return (db.query(ExtractionCandidate)
            .filter(ExtractionCandidate.job_id == job.id,
                    ExtractionCandidate.candidate_kind == "instance")
            .all())


@pytest.fixture()
def _offline(monkeypatch):
    monkeypatch.setattr(settings, "llm_cloud_enabled", False)
    monkeypatch.setattr(settings, "anthropic_api_key", "")


def test_gate_active_drops_unrelated_rows(db, _offline, tmp_path):
    """无 column_mapping：目标类「设备」→ 仅提及设备的行落候选，药品行被门控丢弃。"""
    eng = _StubEngine(_detail(label_zh="设备", name="Equipment"))
    cfg = ExtractionConfig(name="auto-设备", target_class_iri=EQUIP, source_type="word")
    db.add(cfg)
    db.commit()
    db.refresh(cfg)

    job = _job(db)
    asyncio.run(run_extraction_pipeline(
        job, cfg, _mixed_docx(tmp_path / "mix.docx"), eng, db))

    db.refresh(job)
    assert job.status == "reviewing"
    cands = _instances(db, job)
    assert len(cands) == 1
    blob = " ".join(map(str, cands[0].extracted_properties))
    assert "设备" in blob and "活性成分" not in cands[0].extracted_properties


def test_gate_bypassed_when_column_mapping_present(db, _offline, tmp_path):
    """有 column_mapping：分析师已声明映射 → 门控旁路，两行全部落候选（零回归）。"""
    eng = _StubEngine(_detail(label_zh="设备", name="Equipment"))
    cfg = ExtractionConfig(
        name="设备台账", target_class_iri=EQUIP, source_type="word",
        column_mapping={"设备名称": "equipmentName", "活性成分": "activeIngredient"},
    )
    db.add(cfg)
    db.commit()
    db.refresh(cfg)

    job = _job(db)
    asyncio.run(run_extraction_pipeline(
        job, cfg, _mixed_docx(tmp_path / "mix.docx"), eng, db))

    db.refresh(job)
    assert job.status == "reviewing"
    assert len(_instances(db, job)) == 2     # 门控旁路：两行皆保留
